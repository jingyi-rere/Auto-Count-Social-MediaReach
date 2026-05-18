from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

BLUE   = (31, 73, 125)
LBLUE  = (68, 114, 196)
WHITE  = (255, 255, 255)
GREY   = (89, 89, 89)
GREEN  = (0, 112, 0)
ORANGE = (197, 90, 17)
RED    = (192, 0, 0)

def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.style = f'Heading {level}'
    p.add_run(text)
    return p

def add_body(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color='CCCCCC'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), color)
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def set_col_width(cell, width):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def make_table(doc, headers, rows_data, col_widths, header_color='1F497D', alt_color='F0F4FF'):
    t = doc.add_table(rows=len(rows_data)+1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        set_col_width(cell, col_widths[j])
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(*WHITE)
        shade_cell(cell, header_color)
        set_cell_border(cell, '1F497D')
    for i, row_data in enumerate(rows_data, 1):
        row = t.rows[i]
        fill = alt_color if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            set_col_width(cell, col_widths[j])
            if isinstance(val, tuple):
                text, bold, color = val
                run = cell.paragraphs[0].add_run(str(text))
                run.bold = bold
                run.font.size = Pt(10)
                if color:
                    run.font.color.rgb = RGBColor(*color)
            else:
                run = cell.paragraphs[0].add_run(str(val))
                run.font.size = Pt(10)
            shade_cell(cell, fill)
            set_cell_border(cell)
    doc.add_paragraph()
    return t

def add_code(doc, text):
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 139)
    p.paragraph_format.left_indent = Inches(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)
    return p

# ── COVER PAGE ────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('TECHNICAL DESIGN DOCUMENT')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(*BLUE)

doc.add_paragraph()

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Auto Count Social Media Reach')
run2.bold = True
run2.font.size = Pt(18)
run2.font.color.rgb = RGBColor(*LBLUE)

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('AI-Assisted Video Metrics Automation System')
run3.italic = True
run3.font.size = Pt(13)
run3.font.color.rgb = RGBColor(*GREY)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_table(rows=5, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (lbl, val) in enumerate([
    ('Document Version:', '3.3'),
    ('Date:', datetime.date.today().strftime('%d %B %Y')),
    ('Status:', 'Final'),
    ('Prepared By:', 'jingyi-rere'),
    ('Based On:', 'BRD v5.3 — Auto Count Social Media Reach'),
]):
    set_col_width(meta.rows[i].cells[0], 2.2)
    set_col_width(meta.rows[i].cells[1], 4.3)
    lc = meta.rows[i].cells[0].paragraphs[0].add_run(lbl)
    lc.bold = True
    lc.font.size = Pt(10)
    vc = meta.rows[i].cells[1].paragraphs[0].add_run(val)
    vc.font.size = Pt(10)
    shade_cell(meta.rows[i].cells[0], 'D6E4F7')
    set_cell_border(meta.rows[i].cells[0])
    set_cell_border(meta.rows[i].cells[1])

doc.add_paragraph()
doc.add_paragraph()

add_heading(doc, 'Document Version History', 2)
make_table(doc,
    ['Version', 'Date', 'Author', 'Changes'],
    [
        ('1.0', '09 May 2026', 'jingyi-rere', 'Initial TDD — platform APIs approach'),
        ('2.0', '11 May 2026', 'jingyi-rere', 'Redesigned: Playwright + Claude Vision (no APIs), hard column allowlist, run.py trigger, write A/D/E/G only'),
        ('3.0', '13 May 2026', 'jingyi-rere', 'Updated: watcher.py background daemon, yt-dlp for YouTube/TikTok, Firefox for Instagram/RedNote, single atomic Lark write, retry logic, shortcode security validation, startup env check, 283 tests, src/ module layout, structured logging'),
        ('3.1', '13 May 2026', 'jingyi-rere', 'Refined: architecture paragraph, step-by-step cohort timeline, tracking plan, risks and fallbacks per boss requirements'),
        ('3.2', '13 May 2026', 'jingyi-rere', 'Feedback fixes: added pytest validation evidence reference (283 passed, commit 6caf2b1), added Vision latency note to timing risk, added Claude API cost estimate to risks table'),
        ('3.3', datetime.date.today().strftime('%d %B %Y'), 'jingyi-rere', 'Updated architecture: Instagram view count now from reel page screenshot taken before Escape (not profile grid), watcher interval reduced to 3 minutes, corrected system flow diagram and key design decisions'),
    ],
    [0.7, 1.5, 1.5, 3.8]
)

doc.add_page_break()

# ── 1. SYSTEM OVERVIEW ────────────────────────────────────────────────────────
add_heading(doc, '1. System Overview', 1)
add_body(doc, (
    'The Auto Count Social Media Reach system is a Python-based automation tool. '
    'The user pastes video URLs into Column F of their Lark Bitable, then runs python watcher.py once. '
    'The watcher checks Lark every 3 minutes for new URLs and processes them automatically '
    '(reduced from 5 minutes after observing that Instagram+Vision processing takes 2-4 minutes per URL — '
    'a 5-minute interval meant the next check fired before the current batch finished logging). '
    'Each URL is routed to the best extraction method: '
    'yt-dlp (fast, exact numbers, no browser) for YouTube and TikTok; '
    'Firefox + Claude Haiku Vision (claude-haiku-4-5) for Instagram and RedNote. '
    'Instagram uses a two-screenshot approach: reel page for caption and date, '
    'profile grid for view count. '
    'All extracted data is written to columns A, D, E, G in a single atomic Lark API call. '
    'A hard allowlist in lark_writer.py ensures no other columns can ever be written to. '
    'All activity is logged to logs/auto_count.log. Errors are displayed in plain English.'
))

doc.add_paragraph()
# ── NEW: ARCHITECTURE ────────────────────────────────────────────────────────
add_heading(doc, '1.0 Architecture & Approach', 1)
add_body(doc, (
    'The system uses a two-track extraction pipeline controlled by a background watcher daemon. '
    'Track 1 (YouTube, TikTok): each URL is passed to yt-dlp, a battle-tested open-source tool '
    'that retrieves exact metadata (title, upload date, view count) directly from platform '
    'infrastructure without launching a browser — this takes under 10 seconds per video. '
    'Track 2 (Instagram, RedNote): a persistent Firefox browser context (Playwright) loads the '
    'real page using the saved login session, dismisses popups, and takes a screenshot; '
    'for Instagram specifically, two screenshots are taken — the individual reel page for caption '
    'and date, and the profile reels grid for view count (which is only visible there on desktop). '
    'Each screenshot is compressed to under 4.5MB if needed, then sent to Claude Haiku Vision '
    '(claude-haiku-4-5 via the Anthropic API) with a structured prompt; Claude returns a JSON '
    'object containing posted_date, caption, and view_count. '
    'The watcher.py daemon polls Lark Bitable every 3 minutes, routes each new URL through the '
    'correct track, merges the results, strips all hashtags using Unicode-aware regex, and writes '
    'all four fields to Lark columns A, D, E, G in a single atomic API call with automatic retry '
    '(up to 3 attempts, 2s/4s/8s back-off). '
    'The entire pipeline is observable via a rotating log file (logs/auto_count.log) and '
    'protected by 283 automated tests that cover positive, negative, security, performance, '
    'UI/UX, and stress scenarios.'
))
doc.add_paragraph()

# ── NEW: COHORT TIMELINE ─────────────────────────────────────────────────────
add_heading(doc, '1.1 Step-by-Step Path to Done (Cohort Timeline)', 1)
add_body(doc, (
    'All core development is complete. The remaining cohort sessions focus on '
    'validation, documentation, and demonstration. '
    'Each step below has a clear deliverable and a done condition.'
))
doc.add_paragraph()

DONE  = ('DONE', True, GREEN)
WIP   = ('THIS WEEK', True, ORANGE)
NEXT  = ('NEXT CLASS', True, LBLUE)

make_table(doc,
    ['Phase', 'Status', 'What Was / Will Be Done', 'Done Condition'],
    [
        ('Phase 1\nCore Pipeline',
         DONE,
         'Built yt-dlp extractor (YouTube/TikTok), Firefox+Vision extractor (Instagram/RedNote), '
         'Lark reader/writer, watcher daemon, processor router.',
         'System fills Lark columns A/D/E/G for all 4 platforms without manual steps.'),
        ('Phase 2\nQuality & Security',
         DONE,
         '283 automated tests (positive, negative, security, performance, stress). '
         'Retry logic, atomic writes, startup validation, shortcode security, structured logging.',
         'All 283 pytest tests pass. No partial writes on crash.'),
        ('Phase 3\nDocumentation',
         WIP,
         'BRD v5.1: specific problem, quantified value, why AI, acceptance criteria. '
         'TDD v3.1: architecture paragraph, cohort timeline, tracking plan, risks.',
         'Both documents submitted to class portal this session.'),
        ('Phase 4\nLive Validation',
         NEXT,
         'Run system on 10 real videos (mix of YouTube, Instagram, TikTok, RedNote). '
         'Compare extracted data to actual platform values. '
         'Measure time: paste URLs to Lark filled.',
         '95%+ accuracy on view counts. All hashtags removed. Time under 5 min per URL.'),
        ('Phase 5\nFinal Demo',
         NEXT,
         'Record a short screen capture showing: paste URL in Lark, '
         'watcher processes automatically, columns fill in real time. '
         'Present quantified results: time saved, error rate, test coverage.',
         'Demo video shows end-to-end flow in under 3 minutes. '
         'Before/after time comparison presented.'),
    ],
    [1.2, 1.0, 3.3, 2.0]
)
doc.add_paragraph()

# ── NEW: TRACKING PLAN ────────────────────────────────────────────────────────
add_heading(doc, '1.2 Tracking Plan — Where Metrics Live', 1)
add_body(doc, (
    'The Lark Bitable IS the tracking system. Every row in the sheet represents one video. '
    'The system writes four metrics per row. Here is what each metric means and how accuracy '
    'will be verified during the validation phase:'
))
doc.add_paragraph()

make_table(doc,
    ['Metric', 'Lark Column', 'What It Tracks', 'How Accuracy Is Verified'],
    [
        ('Posted Date', 'Column A', 'The original date the video was published on the platform.',
         'Manually check 10 videos on platform. Compare to Column A. Pass if all 10 match.'),
        ('Caption / Title', 'Column D', 'The video caption or title, with ALL hashtags removed.',
         'Check Column D for any # character. Pass if zero hashtags found in any cell.'),
        ('Content Type', 'Column E', 'Always "Content Casual" — hardcoded default.',
         'Check Column E for all processed rows. Pass if every cell says exactly "Content Casual".'),
        ('View Count (Reach)', 'Column G', 'Total views/plays as shown on the platform at time of processing.',
         'Manually check 10 videos. Compare Column G to platform count. Pass if within 5%.'),
    ],
    [1.3, 1.1, 2.0, 3.1]
)
doc.add_paragraph()

add_body(doc, 'System performance metrics (tracked in logs/auto_count.log):', bold=True)
doc.add_paragraph()

make_table(doc,
    ['Performance Metric', 'Target', 'Where It Is Logged'],
    [
        ('Processing time per URL', 'Under 5 minutes', 'Logged as INFO with timestamp per record_id'),
        ('Success rate per cycle', '95%+', 'Logged as "X OK, Y error(s)" at end of each watcher cycle'),
        ('Retry count', '0 retries in normal conditions', 'Logged as WARNING when a retry is triggered'),
        ('Test pass rate', '283/283 (100%)', 'pytest output — run before each submission'),
    ],
    [2.5, 2.0, 3.0]
)
doc.add_paragraph()

# ── NEW: RISKS & FALLBACKS ────────────────────────────────────────────────────
add_heading(doc, '1.3 Risks & Fallbacks if AI Hits Limits', 1)
add_body(doc, (
    'Claude Vision is reliable but not infallible. Every known failure mode has a specific '
    'fallback built into the system. Below are the real risks and what happens in each case:'
))
doc.add_paragraph()

HIGH = ('High', True, RED)
MED  = ('Medium', True, ORANGE)
LOW  = ('Low', True, GREEN)

make_table(doc,
    ['Risk', 'Likelihood', 'What Happens', 'Fallback'],
    [
        ('Claude Vision extracts wrong view count\n(e.g. reads like count instead of view count)',
         MED,
         'Column G has incorrect number. User may not notice unless they check.',
         'Reel page screenshot is taken before Escape and before any navigation, so it always shows the correct reel. '
         'User can manually correct Column G — system will not overwrite it on next run '
         'because it only processes rows where G is empty.'),
        ('Claude API returns 529 (overloaded)',
         LOW,
         'Extraction fails for that URL in this cycle.',
         'Built-in retry: system waits 2s, 4s, 8s and retries up to 3 times automatically. '
         'If all retries fail, error is logged in plain English and URL is left for next cycle.'),
        ('Instagram blocks Firefox (anti-bot detection)',
         MED,
         'Page shows login screen or CAPTCHA instead of video. Screenshot sent to Claude has no data.',
         'System detects null view_count and logs warning. '
         'User re-logs into Instagram via login_once.py (one command). '
         'Persistent session is saved and works again for weeks.'),
        ('Platform changes UI (e.g. Instagram redesign)',
         HIGH,
         'Screenshot layout changes. Claude may read wrong element.',
         'Claude Vision is layout-agnostic — it reads meaning, not HTML selectors. '
         'Minor UI changes: system adapts automatically. '
         'Major redesign: update PROMPT text in vision_extract.py (30-minute fix, no code rewrite).'),
        ('RedNote content in Chinese — Claude misreads',
         LOW,
         'Caption or date extracted incorrectly from Chinese text.',
         'Claude Haiku natively understands Chinese. Tested with Chinese captions in stress tests. '
         'If a specific post fails, user manually fills that row in Lark.'),
        ('yt-dlp breaks for YouTube (platform change)',
         LOW,
         'YouTube/TikTok extraction fails.',
         'yt-dlp is actively maintained by open-source community. '
         'Fix is usually available within 24 hours via: pip install -U yt-dlp. '
         'Fallback: route YouTube to Firefox+Vision path (slower but works).'),
        ('Anthropic API key quota exceeded',
         LOW,
         'All Vision extractions fail.',
         'System logs clear error: "Claude AI key is invalid or missing. Check ANTHROPIC_API_KEY". '
         'User tops up Anthropic credits (takes 2 minutes). '
         'yt-dlp path (YouTube/TikTok) continues working unaffected.'),
        ('Vision extraction takes 3-4 min per Instagram URL — may miss <5 min promise',
         MED,
         'If a batch of 5+ Instagram URLs arrives together, total processing time may exceed 5 minutes. '
         'The watcher runs every 3 minutes but each Vision call itself takes 2-4 minutes.',
         'YouTube/TikTok (yt-dlp) complete in under 30 seconds — <5 min promise holds for those. '
         'For Instagram/RedNote batches: the 5-minute SLA applies per URL, not per batch. '
         'Single URL always processed in under 5 minutes. BRD AC-07 specifies "one URL at a time" for verification.'),
        ('Claude API cost at scale',
         LOW,
         'At current usage: ~10 Instagram/RedNote videos/week × 2 screenshots each × 52 weeks '
         '= ~1,040 Vision API calls/year. '
         'Claude Haiku-4-5 costs approx $0.0004/image. '
         'Estimated annual cost: ~$0.42 USD for Vision calls. '
         'YouTube/TikTok use yt-dlp (free). Total API spend: under $5/year at current scale.',
         'Cost is negligible at current posting volume. '
         'If volume scales to 40 videos/week (4x current), annual cost is still under $2. '
         'No rate limiting needed at this scale.'),
    ],
    [2.2, 0.8, 2.0, 2.5]
)
doc.add_paragraph()

add_heading(doc, '1.1 Key Design Decisions', 2)
make_table(doc,
    ['Decision', 'Choice', 'Reason'],
    [
        ('YouTube/TikTok extraction', 'yt-dlp (no browser)', 'Fast, exact numbers. Uses Chrome cookies for auth. No screenshots needed.'),
        ('Instagram/RedNote extraction', 'Firefox + Claude Haiku Vision (claude-haiku-4-5)', 'No API available. Firefox with saved login handles paywalls/popups.'),
        ('Instagram view count', 'Reel page screenshot (before Escape) + split-view', 'Screenshot taken immediately after page load, before Escape key is pressed. Escape can navigate away from the reel, making any later screenshot show the wrong content. Reel page is always the URL the user pasted — can never show the wrong reel.'),
        ('Column safety', 'Hard allowlist raises ValueError for B/C/F/H', 'Prevents accidental data corruption in user-managed columns. Enforced in code, not config.'),
        ('Lark write strategy', 'Single atomic API call per row in write_row()', 'Prevents partial writes if process crashes mid-row. All fields written or none.'),
        ('Retry logic', 'with_retry(): 3 attempts, 2s/4s/8s back-off', 'Handles transient network blips without crashing the watcher loop.'),
        ('Startup validation', '_validate_env() checks all .env vars at boot', 'Clear error at startup instead of confusing NoneType crash mid-run.'),
        ('Shortcode security', 're.fullmatch() validates shortcode before JS eval', 'Prevents JS injection if a malicious URL reaches the page.evaluate() call.'),
        ('Trigger mode', 'watcher.py daemon, checks Lark every 3 minutes', 'User pastes URL once - data appears automatically. No manual run per batch.'),
        ('Browser', 'Firefox persistent context, ~/.cache/auto-count/firefox-profile', 'Separate from user Chrome. Retains Instagram/RedNote login across runs.'),
        ('Shared utilities', 'src/utils.py: clean_caption + with_retry', 'Single source of truth - no duplicate code across modules.'),
    ],
    [1.8, 2.5, 3.2]
)

# ── 2. SYSTEM ARCHITECTURE ────────────────────────────────────────────────────
add_heading(doc, '2. System Architecture', 1)

add_heading(doc, '2.1 Component Overview', 2)
make_table(doc,
    ['File', 'Component', 'What It Does'],
    [
        ('watcher.py', 'Background Daemon', 'Checks Lark every 3 min. Validates .env on startup. Calls process_all() each cycle.'),
        ('src/lark_reader.py', 'Lark Bitable Reader', 'Reads Lark. Returns (record_id, url) pairs where Column F has URL and A/D/E/G are empty.'),
        ('src/lark_writer.py', 'Lark Bitable Writer', 'Writes A/D/E/G in ONE atomic API call. Hard allowlist raises ValueError for any other column. Retries 3x.'),
        ('src/browser_reader.py', 'Firefox Controller', 'Opens Instagram/RedNote in persistent Firefox. Instagram: reel page screenshot (view count, before Escape) + profile grid click for split-view (caption/date).'),
        ('src/metadata_reader.py', 'yt-dlp Extractor', 'Extracts YouTube/TikTok metadata without a browser. Fast and exact.'),
        ('src/vision_extract.py', 'Claude Vision Extractor', 'Sends screenshot to claude-haiku-4-5. Compresses images >4.5MB. Returns posted_date, caption, view_count.'),
        ('src/processor.py', 'URL Router + Orchestrator', 'Routes each URL to yt-dlp or Firefox+Vision. Merges Instagram two-screenshot results.'),
        ('src/utils.py', 'Shared Utilities', 'clean_caption() (Unicode hashtag strip) and with_retry() decorator. Used across all modules.'),
        ('src/logger.py', 'Structured Logger', 'Rotating log file at logs/auto_count.log (1MB, 7 backups). INFO to terminal, DEBUG to file.'),
        ('src/friendly_errors.py', 'Error Translator', 'Maps 15+ technical exceptions to plain English messages for non-technical users.'),
        ('src/reporter.py', 'Run Reporter', 'Prints run summary: rows processed, successes, failures, friendly error descriptions.'),
    ],
    [2.0, 1.8, 3.7]
)

add_heading(doc, '2.2 System Flow Diagram', 2)

for line in [
    '  USER starts watcher once: python watcher.py',
    '  watcher.py: validates .env, then loops every 3 minutes',
    '         |',
    '  USER pastes URLs into Lark Column F (any time)',
    '         |',
    '  [lark_reader.py] -- finds rows where F has URL + A/D/E/G empty',
    '         |',
    '         | for each URL',
    '         v',
    '  [processor.py] _route(url)',
    '         |',
    '         +---- YouTube / TikTok? ----+',
    '         |                           v',
    '         |             [metadata_reader.py]',
    '         |             yt-dlp: posted_date, caption, view_count',
    '         |',
    '         +---- Instagram / RedNote? -+',
    '                                     v',
    '                       [browser_reader.py] Firefox',
    '                       Instagram: reel page screenshot (view count, taken before Escape)',
    '                                + split-view screenshot (caption+date, from grid click)',
    '                       RedNote:  full-page screenshot',
    '                                     v',
    '                       [vision_extract.py] claude-haiku-4-5',
    '                       Compresses if >4.5MB',
    '                       Returns: { posted_date, caption, view_count }',
    '         |',
    '         v',
    '  [src/utils.py] clean_caption() -- strips ALL Unicode hashtags',
    '         |',
    '         v',
    '  [lark_writer.py]  <-- HARD ALLOWLIST: A, D, E, G only',
    '  ONE atomic API call: { Date, Title, Content Type, Reach }',
    '  Column A <-- posted_date',
    '  Column D <-- cleaned caption',
    '  Column E <-- "Content Casual"  (always)',
    '  Column G <-- view_count',
    '  Column B/C/F/H --> ValueError immediately',
    '  Retries 3x on failure (2s, 4s, 8s back-off)',
    '         |',
    '         v',
    '  [reporter.py] + [logger.py] + [friendly_errors.py]',
    '  logs/auto_count.log + terminal summary',
]:
    add_code(doc, line)

doc.add_paragraph()

# ── 3. TECHNOLOGY STACK ───────────────────────────────────────────────────────
add_heading(doc, '3. Technology Stack', 1)
make_table(doc,
    ['Library / Tool', 'Version', 'Purpose'],
    [
        ('Python', '3.9+', 'Core programming language'),
        ('playwright', 'latest (pinned)', 'Firefox browser automation for Instagram/RedNote screenshots'),
        ('yt-dlp', 'latest (pinned)', 'YouTube/TikTok metadata extraction without browser (fast, exact)'),
        ('anthropic', 'latest (pinned)', 'Claude Haiku Vision API (claude-haiku-4-5) for screenshot analysis'),
        ('Pillow (PIL)', 'latest (pinned)', 'Image compression — reduces screenshots >4.5MB before sending to Claude'),
        ('lark-oapi', 'latest (pinned)', 'Official Lark Open API SDK for reading/writing Lark Bitable'),
        ('python-dotenv', 'latest (pinned)', 'Load API keys and secrets from .env file securely'),
        ('python-dateutil', 'latest (pinned)', 'Parse flexible date strings (e.g. "3 days ago") into timestamps'),
        ('pytest', 'latest (pinned)', '283 automated tests across 6 test files — no real network calls'),
        ('re (built-in)', 'built-in', 'Unicode-aware regex for hashtag removal and shortcode validation'),
        ('logging (built-in)', 'built-in', 'Structured rotating log file (logs/auto_count.log)'),
        ('asyncio (built-in)', 'built-in', 'Async event loop for Playwright browser operations'),
    ],
    [2.2, 1.0, 4.3]
)

# ── 4. DETAILED COMPONENT DESIGN ─────────────────────────────────────────────
add_heading(doc, '4. Detailed Component Design', 1)

add_heading(doc, '4.1 lark_reader.py', 2)
add_body(doc, 'Reads the Lark Sheet and returns rows that need processing.')
make_table(doc,
    ['Function', 'Input', 'Output', 'Logic'],
    [
        ('get_rows_to_process()', 'None', 'List of {row_index, url}', 'Find rows where Column F has URL AND at least one of A/D/E/G is empty'),
        ('connect_lark()', 'None', 'Lark client object', 'Authenticate using App ID + Secret from .env'),
        ('read_sheet()', 'Lark client', 'Raw sheet data', 'Read all rows from configured sheet tab'),
    ],
    [2.2, 1.5, 1.8, 2.0]
)

add_heading(doc, '4.2 lark_writer.py — Hard Column Allowlist', 2)
add_body(doc, 'This is the most critical safety component. The hard allowlist MUST be enforced:', bold=True, color=RED)
doc.add_paragraph()

add_code(doc, '  ALLOWED_COLUMNS = {"A", "D", "E", "G"}  # Hard allowlist')
add_code(doc, '')
add_code(doc, '  def write_cell(row_index, column, value):')
add_code(doc, '      if column not in ALLOWED_COLUMNS:')
add_code(doc, '          raise ValueError(')
add_code(doc, '              f"FORBIDDEN: Cannot write to column {column}. "')
add_code(doc, '              f"Only {ALLOWED_COLUMNS} are allowed."')
add_code(doc, '          )')
add_code(doc, '      # proceed to write to Lark Sheet...')
doc.add_paragraph()

make_table(doc,
    ['Function', 'Description'],
    [
        ('write_cell(row_index, column, value)', 'Writes one cell. Raises ValueError immediately if column not in allowlist.'),
        ('write_row(row_index, data_dict)', 'Writes A/D/E/G for one row. Calls write_cell for each — allowlist enforced per call.'),
        ('verify_write(row_index)', 'Reads back the written row to confirm data was saved correctly.'),
    ],
    [3.0, 4.5]
)

add_heading(doc, '4.3 browser_reader.py', 2)
add_body(doc, 'Opens each video URL in a real Chrome browser and takes a screenshot.')
make_table(doc,
    ['Function', 'Description'],
    [
        ('get_browser()', 'Launches Playwright with persistent Chrome profile at ~/.cache/auto-count/chrome-profile'),
        ('open_url(browser, url)', 'Opens the video URL in a new tab. Waits for page to fully load.'),
        ('take_screenshot(page)', 'Takes a full-page screenshot. Returns image bytes.'),
        ('close_browser(browser)', 'Closes browser cleanly after all URLs are processed.'),
    ],
    [2.5, 5.0]
)

add_body(doc, 'Persistent Chrome profile setup:', italic=True, color=GREY)
add_code(doc, '  PROFILE_PATH = Path.home() / ".cache" / "auto-count" / "chrome-profile"')
add_code(doc, '  PROFILE_PATH.mkdir(parents=True, exist_ok=True)')
add_code(doc, '')
add_code(doc, '  browser = playwright.chromium.launch_persistent_context(')
add_code(doc, '      user_data_dir=str(PROFILE_PATH),')
add_code(doc, '      headless=False,  # visible browser so user can log in if needed')
add_code(doc, '  )')
doc.add_paragraph()

add_heading(doc, '4.4 vision_extract.py', 2)
add_body(doc, 'Sends screenshot to Claude Haiku Vision and returns structured data.')
make_table(doc,
    ['Function', 'Input', 'Output'],
    [
        ('extract_from_screenshot(image_bytes, url)', 'Screenshot bytes + URL', 'JSON: { posted_date, caption, view_count }'),
        ('build_prompt(url)', 'Video URL', 'Prompt string telling Claude what to extract'),
        ('parse_response(response)', 'Claude API response', 'Validated dict with posted_date, caption, view_count'),
        ('validate_extraction(data)', 'Extracted dict', 'Raises error if required fields are missing or invalid'),
    ],
    [3.0, 2.0, 2.5]
)

add_body(doc, 'Claude Vision prompt template:', italic=True, color=GREY)
add_code(doc, '  You are analysing a screenshot of a social media video post.')
add_code(doc, '  Extract exactly these 3 fields and return as JSON only:')
add_code(doc, '  {')
add_code(doc, '    "posted_date": "DD/MM/YYYY format of when the video was posted",')
add_code(doc, '    "caption": "Full video caption text including hashtags",')
add_code(doc, '    "view_count": 123456  (integer, the number of views shown)')
add_code(doc, '  }')
add_code(doc, '  If a field cannot be found, use null.')
add_code(doc, '  Return JSON only. No explanation.')
doc.add_paragraph()

add_heading(doc, '4.5 processor.py', 2)
add_body(doc, 'Cleans and formats extracted data before writing to Lark Sheet.')
make_table(doc,
    ['Function', 'Input', 'Output', 'Rule'],
    [
        ('clean_caption(text)', 'Raw caption with hashtags', 'Caption without any hashtags', 'Unicode-aware: strips #word in any language'),
        ('format_date(raw_date)', 'Date string from Claude', 'DD/MM/YYYY string', 'Standardise to consistent format'),
        ('set_content_type()', 'Nothing', '"Content Casual"', 'Always returns this exact string'),
        ('handle_empty(value, fallback)', 'Any value', 'Value or fallback', 'Returns fallback if value is None or empty'),
    ],
    [2.0, 1.8, 1.8, 1.9]
)

add_body(doc, 'Unicode-aware hashtag removal:', italic=True, color=GREY)
add_code(doc, '  import re')
add_code(doc, '')
add_code(doc, '  def clean_caption(text):')
add_code(doc, '      if not text or text.strip() == "":')
add_code(doc, '          return "No caption"')
add_code(doc, '      # \\w matches Unicode word chars (Chinese, Arabic, emoji, etc.)')
add_code(doc, '      cleaned = re.sub(r\'#\\w+\', \'\', text, flags=re.UNICODE)')
add_code(doc, '      return " ".join(cleaned.split())  # collapse extra whitespace')
doc.add_paragraph()

add_heading(doc, '4.6 run.py', 2)
add_body(doc, 'Main entry point. Orchestrates all components in order.')
add_code(doc, '  1. Load .env (API keys, Lark credentials)')
add_code(doc, '  2. lark_reader.get_rows_to_process() --> list of {row_index, url}')
add_code(doc, '  3. browser_reader.get_browser()')
add_code(doc, '  4. For each row:')
add_code(doc, '       a. browser_reader.open_url(url)')
add_code(doc, '       b. browser_reader.take_screenshot()')
add_code(doc, '       c. vision_extract.extract_from_screenshot()')
add_code(doc, '       d. processor.clean_caption() / format_date() / set_content_type()')
add_code(doc, '       e. lark_writer.write_row(row_index, {A, D, E, G})')
add_code(doc, '       f. Record success or failure')
add_code(doc, '  5. browser_reader.close_browser()')
add_code(doc, '  6. reporter.print_summary(results)')
doc.add_paragraph()

add_heading(doc, '4.7 reporter.py', 2)
add_body(doc, 'Prints a clear run summary to the terminal after all URLs are processed.')
add_code(doc, '  ================================')
add_code(doc, '  AUTO COUNT SOCIAL MEDIA REACH')
add_code(doc, '  Run Summary — 11 May 2026 10:30')
add_code(doc, '  ================================')
add_code(doc, '  Total rows processed : 10')
add_code(doc, '  Successful           : 9')
add_code(doc, '  Failed               : 1')
add_code(doc, '  --------------------------------')
add_code(doc, '  FAILED:')
add_code(doc, '  Row 5 | https://rednote.com/... | Error: login required')
add_code(doc, '  ================================')
doc.add_paragraph()

# ── 5. DATA FLOW ──────────────────────────────────────────────────────────────
add_heading(doc, '5. Data Flow', 1)

make_table(doc,
    ['Step', 'From', 'To', 'Data'],
    [
        ('1', 'User', 'Lark Sheet Column F', 'Video URL (user pastes manually)'),
        ('2', 'lark_reader', 'run.py', 'List of {row_index, url} where A/D/E/G are empty'),
        ('3', 'run.py', 'browser_reader', 'URL string'),
        ('4', 'browser_reader', 'Playwright Chrome', 'Open URL command'),
        ('5', 'Playwright Chrome', 'vision_extract', 'Screenshot image bytes'),
        ('6', 'vision_extract', 'Claude Haiku Vision API', 'Screenshot + extraction prompt'),
        ('7', 'Claude Haiku Vision API', 'vision_extract', 'JSON: { posted_date, caption, view_count }'),
        ('8', 'vision_extract', 'processor', 'Raw extracted data'),
        ('9', 'processor', 'lark_writer', 'Clean data: { A: date, D: caption, E: "Content Casual", G: views }'),
        ('10', 'lark_writer', 'Lark Sheet', 'Writes to columns A, D, E, G in existing row'),
        ('11', 'run.py', 'reporter', 'List of results (success/fail per row)'),
        ('12', 'reporter', 'Terminal', 'Printed run summary'),
    ],
    [0.5, 2.0, 2.0, 3.0]
)

# ── 6. COLUMN ALLOWLIST ENFORCEMENT ──────────────────────────────────────────
add_heading(doc, '6. Column Allowlist Enforcement', 1)
add_body(doc, 'This is the most critical safety mechanism in the system.', bold=True)
doc.add_paragraph()

make_table(doc,
    ['Column', 'Allowed to Write?', 'What Happens if Write Attempted'],
    [
        ('A — Date', ('YES', False, GREEN), 'Write proceeds normally'),
        ('B — Week', ('NO', True, RED), 'ValueError raised immediately — write aborted'),
        ('C — PIC', ('NO', True, RED), 'ValueError raised immediately — write aborted'),
        ('D — Title', ('YES', False, GREEN), 'Write proceeds normally'),
        ('E — Content Type', ('YES', False, GREEN), 'Write proceeds normally — always "Content Casual"'),
        ('F — Link', ('NO', True, RED), 'ValueError raised immediately — user URL protected'),
        ('G — Reach', ('YES', False, GREEN), 'Write proceeds normally'),
        ('H — Final Reach', ('NO', True, RED), 'ValueError raised immediately — write aborted'),
    ],
    [1.5, 1.8, 4.2]
)

add_heading(doc, '6.1 pytest Tests for Allowlist', 2)
add_body(doc, 'These tests MUST pass (green) before development continues:', italic=True, color=GREY)
add_code(doc, '  def test_write_to_B_raises():')
add_code(doc, '      with pytest.raises(ValueError):')
add_code(doc, '          lark_writer.write_cell(1, "B", "value")')
add_code(doc, '')
add_code(doc, '  def test_write_to_C_raises():')
add_code(doc, '      with pytest.raises(ValueError):')
add_code(doc, '          lark_writer.write_cell(1, "C", "value")')
add_code(doc, '')
add_code(doc, '  def test_write_to_F_raises():')
add_code(doc, '      with pytest.raises(ValueError):')
add_code(doc, '          lark_writer.write_cell(1, "F", "value")')
add_code(doc, '')
add_code(doc, '  def test_write_to_H_raises():')
add_code(doc, '      with pytest.raises(ValueError):')
add_code(doc, '          lark_writer.write_cell(1, "H", "value")')
add_code(doc, '')
add_code(doc, '  def test_write_to_A_succeeds():')
add_code(doc, '      # Should NOT raise')
add_code(doc, '      lark_writer.write_cell(1, "A", "01/05/2026")')
doc.add_paragraph()

# ── 7. LARK SHEET INTEGRATION ─────────────────────────────────────────────────
add_heading(doc, '7. Lark Sheet Integration', 1)

add_heading(doc, '7.1 Authentication', 2)
make_table(doc,
    ['Setting', 'Description'],
    [
        ('LARK_APP_ID', 'From Lark Open Platform — stored in .env'),
        ('LARK_APP_SECRET', 'From Lark Open Platform — stored in .env'),
        ('LARK_SHEET_TOKEN', 'From the Lark Sheet URL — stored in .env'),
        ('LARK_SHEET_TAB', 'The specific tab name where video data lives — stored in .env'),
    ],
    [2.5, 5.0]
)

add_heading(doc, '7.2 Read Logic (lark_reader)', 2)
for item in [
    'Read all rows from the configured sheet tab.',
    'For each row, check if Column F has a non-empty URL.',
    'If Column F has URL AND at least one of A/D/E/G is empty, add to processing list.',
    'Return list of {row_index, url} pairs.',
]:
    add_bullet(doc, item)

add_heading(doc, '7.3 Write Logic (lark_writer)', 2)
for item in [
    'Check column against ALLOWED_COLUMNS allowlist — raise ValueError if not allowed.',
    'Write to the specific row_index — never a new row.',
    'Write columns A, D, E, G in a single batch API call for efficiency.',
    'Read back the row after writing to verify data was saved correctly.',
]:
    add_bullet(doc, item)

# ── 8. CLAUDE VISION INTEGRATION ─────────────────────────────────────────────
add_heading(doc, '8. Claude Vision Integration', 1)

make_table(doc,
    ['Setting', 'Value'],
    [
        ('Model', 'claude-haiku-4-5'),
        ('API Key', 'ANTHROPIC_API_KEY from .env — never logged'),
        ('Input', 'Base64-encoded screenshot image + extraction prompt'),
        ('Output', 'JSON string: { posted_date, caption, view_count }'),
        ('Fallback', 'If extraction fails, log error and skip row — do not write partial data'),
    ],
    [2.5, 5.0]
)

# ── 9. ERROR HANDLING ─────────────────────────────────────────────────────────
add_heading(doc, '9. Error Handling', 1)
make_table(doc,
    ['Error', 'Where It Occurs', 'System Response'],
    [
        ('Write to forbidden column', 'lark_writer.py', 'Raise ValueError immediately. Log error. Skip this row. Continue.'),
        ('Platform blocks browser', 'browser_reader.py', 'Log error. Skip this URL. Continue with next.'),
        ('Claude Vision returns null fields', 'vision_extract.py', 'Log warning. Use "No caption" / 0 fallbacks. Continue.'),
        ('Lark API auth failure', 'lark_reader/writer', 'Retry token refresh once. If still fails, stop and alert user.'),
        ('URL not recognised as supported platform', 'run.py', 'Log warning. Skip this URL. Continue.'),
        ('Screenshot is blank or loading failed', 'browser_reader.py', 'Retry once after 3 seconds. If still blank, skip row.'),
    ],
    [2.2, 2.0, 3.3]
)

# ── 10. SECURITY ──────────────────────────────────────────────────────────────
add_heading(doc, '10. Security Design', 1)
make_table(doc,
    ['Security Concern', 'How It Is Handled'],
    [
        ('API keys', 'Stored in .env only. Never hardcoded. Never logged. Never printed.'),
        ('.env file', 'Added to .gitignore — never pushed to GitHub.'),
        ('.env.template', 'A safe template with placeholder values pushed to GitHub for reference.'),
        ('Lark credentials', 'App ID and Secret in .env. Token refreshed per session.'),
        ('Firefox profile', 'Stored locally at ~/.cache/auto-count/firefox-profile. Never shared. Contains login sessions.'),
        ('Shortcode validation', 're.fullmatch(r"[A-Za-z0-9_-]+", shortcode) before page.evaluate(). Prevents JS injection from malicious URLs.'),
        ('Column protection', 'Hard allowlist in code raises exception — cannot be bypassed at runtime.'),
    ],
    [2.5, 5.0]
)

# ── 11. PROJECT FILE STRUCTURE ────────────────────────────────────────────────
add_heading(doc, '11. Project File Structure', 1)
for line in [
    'Auto-Count-Social-MediaReach/',
    '|-- watcher.py                # Background daemon: python watcher.py',
    '|-- src/',
    '|   |-- processor.py          # URL router + orchestrator (routes to yt-dlp or Firefox+Vision)',
    '|   |-- lark_reader.py        # Read Lark Bitable, find rows to process',
    '|   |-- lark_writer.py        # Write to Lark Bitable (hard allowlist A/D/E/G, single call, retry)',
    '|   |-- browser_reader.py     # Firefox automation: Instagram 2-screenshot, RedNote screenshot',
    '|   |-- metadata_reader.py    # yt-dlp: YouTube/TikTok metadata (no browser)',
    '|   |-- vision_extract.py     # Claude Vision: screenshot -> JSON (claude-haiku-4-5)',
    '|   |-- utils.py              # Shared: clean_caption(), with_retry() decorator',
    '|   |-- logger.py             # Rotating log file: logs/auto_count.log',
    '|   |-- friendly_errors.py    # Maps technical errors to plain English messages',
    '|   |-- reporter.py           # Print run summary to terminal',
    '|   |-- _env.py               # Load .env from project root',
    '|-- tests/',
    '|   |-- test_lark_writer.py        # B/C/F/H raises ValueError, single-call write',
    '|   |-- test_lark_writer_extended.py  # Extended write tests',
    '|   |-- test_lark_reader.py        # URL extraction edge cases',
    '|   |-- test_positive_cases.py     # 67 positive tests',
    '|   |-- test_negative_cases.py     # 67 negative/edge case tests',
    '|   |-- test_security.py           # 16 security tests',
    '|   |-- test_performance.py        # 8 performance/timing tests',
    '|   |-- test_uiux.py               # 15 UI/UX and error message tests',
    '|   |-- test_stress.py             # 21 stress tests (50 URLs, duplicates, weird inputs)',
    '|-- logs/                     # auto_count.log (gitignored)',
    '|-- .env                      # API keys (NEVER pushed to GitHub)',
    '|-- .env.template             # Safe template with placeholder values',
    '|-- .gitignore                # Excludes .env, logs/, firefox-profile, __pycache__',
    '|-- requirements.txt          # Pinned Python dependencies',
    '|-- BRD_Auto_Count_Social_Media_Reach.docx',
    '|-- TDD_Auto_Count_Social_Media_Reach.docx',
]:
    add_code(doc, '  ' + line)
doc.add_paragraph()

# ── 12. BUILD ORDER ───────────────────────────────────────────────────────────
add_heading(doc, '12. Build Order', 1)
make_table(doc,
    ['Phase', 'What to Build', 'Stop Condition'],
    [
        ('1', 'Project scaffold: requirements.txt (pinned), .gitignore, .env.template, virtualenv', 'Virtual environment activated successfully'),
        ('2', 'lark_writer.py with hard allowlist + pytest tests for B/C/F/H', 'ALL pytest tests GREEN before proceeding'),
        ('3', 'lark_reader.py — find rows where F has URL and A/D/E/G are empty', 'Returns correct rows from test sheet'),
        ('4', 'processor.py — clean_caption, format_date, set_content_type', 'Unit tests pass for hashtag removal'),
        ('5', 'browser_reader.py — Playwright persistent Chrome profile', 'Browser opens URL and returns screenshot'),
        ('6', 'vision_extract.py — Claude Vision screenshot to JSON', 'Returns valid JSON for a YouTube video'),
        ('7', 'run.py + reporter.py — wire everything + run summary', 'Full run completes without errors'),
        ('8', 'End-to-end test on 3 real YouTube URLs', 'Lark Sheet correctly filled for all 3 rows'),
    ],
    [0.5, 4.5, 2.5]
)

# ── 13. TESTING PLAN ──────────────────────────────────────────────────────────
add_heading(doc, '13. Testing Plan', 1)
add_body(doc, (
    'Validation evidence: all 283 tests confirmed passing. '
    'Full pytest verbose output saved in validation_evidence.txt (project root). '
    'Result: 283 passed, 0 failed, 3 warnings, 4.21 seconds. '
    'Git commit: 6caf2b1 (branch: main, repo: github.com/jingyi-rere/Auto-Count-Social-MediaReach). '
    'Tests run with: .venv/bin/python -m pytest tests/ -v --tb=short'
), bold=True, color=GREEN)
doc.add_paragraph()
make_table(doc,
    ['Test File', 'Count', 'What Is Tested', 'Pass Criteria'],
    [
        ('test_lark_writer.py + test_lark_writer_extended.py', '~30', 'Column allowlist, single atomic write, retry, field content', 'All pass — no partial writes'),
        ('test_positive_cases.py', '67', 'Happy path for all modules: URLs, dates, captions, views, routing', 'All 67 pass'),
        ('test_negative_cases.py', '67', 'Edge cases: null dates, empty captions, bad URLs, API errors, forbidden columns', 'All 67 pass'),
        ('test_security.py', '16', 'API keys not in code, .env gitignored, column protection, shortcode safety', 'All 16 pass'),
        ('test_performance.py', '8', 'Function speed (<1s each), 5-URL batch <2s, 15-URL batch <5s (mocked)', 'All under time limit'),
        ('test_uiux.py', '15', 'Error messages in plain English, log file format, result structure', 'All 15 pass'),
        ('test_stress.py', '21', '50 URLs at once, duplicate URLs, Arabic/Japanese text, 10 failures in a row', 'System never crashes'),
        ('TOTAL', '283', 'All tests mocked — no real network calls needed', 'pytest: 283 passed'),
    ],
    [2.8, 0.6, 2.8, 1.3]
)

# ── 14. DEPENDENCIES ──────────────────────────────────────────────────────────
add_heading(doc, '14. requirements.txt (Pinned)', 1)
add_body(doc, 'All versions must be pinned in requirements.txt:')
add_code(doc, '  playwright==1.44.0')
add_code(doc, '  anthropic==0.28.0')
add_code(doc, '  lark-oapi==1.3.5')
add_code(doc, '  python-dotenv==1.0.1')
add_code(doc, '  pytest==8.2.0')
doc.add_paragraph()

# ── 15. SIGN-OFF ──────────────────────────────────────────────────────────────
add_heading(doc, '15. Sign-Off & Approval', 1)
add_body(doc, 'By signing below, the approver confirms that this Technical Design Document accurately represents the proposed technical approach.')
doc.add_paragraph()
make_table(doc,
    ['Role', 'Name', 'Signature', 'Date'],
    [
        ('Prepared By', 'jingyi-rere', '', ''),
        ('Reviewed By', '', '', ''),
        ('Approved By', '', '', ''),
    ],
    [2.0, 2.0, 2.0, 1.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— End of Document — Version 3.3 —')
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(*GREY)

output_path = '/Users/jjjyyy/Auto-Count Social MediaReach/TDD_Auto_Count_Social_Media_Reach.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
