from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

BLUE   = (31, 73, 125)
LBLUE  = (68, 114, 196)
WHITE  = (255, 255, 255)
GREY   = (89, 89, 89)
GREEN  = (0, 112, 0)
ORANGE = (197, 90, 17)
RED    = (192, 0, 0)

def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.style = f'Heading {level}'
    p.add_run(text)
    return p

def add_body(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color='CCCCCC'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), color)
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def set_col_width(cell, width):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def make_table(doc, headers, rows_data, col_widths, header_color='1F497D', alt_color='F0F4FF'):
    t = doc.add_table(rows=len(rows_data)+1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        set_col_width(cell, col_widths[j])
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(*WHITE)
        shade_cell(cell, header_color)
        set_cell_border(cell, '1F497D')
    for i, row_data in enumerate(rows_data, 1):
        row = t.rows[i]
        fill = alt_color if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            set_col_width(cell, col_widths[j])
            if isinstance(val, tuple):
                text, bold, color = val
                run = cell.paragraphs[0].add_run(str(text))
                run.bold = bold
                run.font.size = Pt(10)
                if color:
                    run.font.color.rgb = RGBColor(*color)
            else:
                run = cell.paragraphs[0].add_run(str(val))
                run.font.size = Pt(10)
            shade_cell(cell, fill)
            set_cell_border(cell)
    doc.add_paragraph()
    return t

def add_code(doc, text):
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 139)
    p.paragraph_format.left_indent = Inches(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)
    return p

# ── COVER PAGE ────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('TECHNICAL DESIGN DOCUMENT')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(*BLUE)

doc.add_paragraph()

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Auto Count Social Media Reach')
run2.bold = True
run2.font.size = Pt(18)
run2.font.color.rgb = RGBColor(*LBLUE)

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('AI-Assisted Video Metrics Automation System')
run3.italic = True
run3.font.size = Pt(13)
run3.font.color.rgb = RGBColor(*GREY)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_table(rows=5, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (lbl, val) in enumerate([
    ('Document Version:', '2.0'),
    ('Date:', datetime.date.today().strftime('%d %B %Y')),
    ('Status:', 'Final'),
    ('Prepared By:', 'jingyi-rere'),
    ('Based On:', 'BRD v4.0 — Auto Count Social Media Reach'),
]):
    set_col_width(meta.rows[i].cells[0], 2.2)
    set_col_width(meta.rows[i].cells[1], 4.3)
    lc = meta.rows[i].cells[0].paragraphs[0].add_run(lbl)
    lc.bold = True
    lc.font.size = Pt(10)
    vc = meta.rows[i].cells[1].paragraphs[0].add_run(val)
    vc.font.size = Pt(10)
    shade_cell(meta.rows[i].cells[0], 'D6E4F7')
    set_cell_border(meta.rows[i].cells[0])
    set_cell_border(meta.rows[i].cells[1])

doc.add_paragraph()
doc.add_paragraph()

add_heading(doc, 'Document Version History', 2)
make_table(doc,
    ['Version', 'Date', 'Author', 'Changes'],
    [
        ('1.0', '09 May 2026', 'jingyi-rere', 'Initial TDD — platform APIs approach'),
        ('2.0', datetime.date.today().strftime('%d %B %Y'), 'jingyi-rere', 'Redesigned: Playwright + Claude Vision (no APIs), hard column allowlist, run.py trigger, write A/D/E/G only'),
    ],
    [0.7, 1.5, 1.5, 3.8]
)

doc.add_page_break()

# ── 1. SYSTEM OVERVIEW ────────────────────────────────────────────────────────
add_heading(doc, '1. System Overview', 1)
add_body(doc, (
    'The Auto Count Social Media Reach system is a Python-based automation tool. '
    'The user pastes video URLs into Column F of their Lark Sheet, then runs python run.py. '
    'The system reads those URLs, opens each video in a real Chrome browser using Playwright, '
    'takes a screenshot, and sends it to Claude Haiku Vision to extract the posted date, '
    'caption, and view count. The data is then written to columns A, D, E, G of the Lark Sheet. '
    'A hard allowlist in lark_writer.py enforces that no other columns can ever be written to.'
))

doc.add_paragraph()
add_heading(doc, '1.1 Key Design Decisions', 2)
make_table(doc,
    ['Decision', 'Choice', 'Reason'],
    [
        ('Data extraction method', 'Playwright + Claude Vision (NO platform APIs)', 'Works on all 7 platforms including RedNote which has no API. More reliable and future-proof.'),
        ('Column safety', 'Hard allowlist raises exception for B/C/F/H', 'Prevents any accidental data corruption in user-managed columns.'),
        ('Column F', 'Read only — never write', 'User owns this column. System reads URL from here but never modifies it.'),
        ('Row management', 'Write into existing rows only', 'User controls row structure. System fills in the blanks only.'),
        ('Trigger', 'python run.py — no UI', 'Simple, fast, no dependencies on web frameworks.'),
        ('AI model', 'claude-haiku-4-5-20251001 for vision', 'Fast and cost-effective for screenshot analysis.'),
        ('Browser profile', 'Persistent Chrome at ~/.cache/auto-count/chrome-profile', 'Retains login sessions across runs — no re-login needed.'),
    ],
    [1.8, 2.2, 3.5]
)

# ── 2. SYSTEM ARCHITECTURE ────────────────────────────────────────────────────
add_heading(doc, '2. System Architecture', 1)

add_heading(doc, '2.1 Component Overview', 2)
make_table(doc,
    ['File', 'Component', 'What It Does'],
    [
        ('run.py', 'Main Entry Point', 'Orchestrates all components. User runs this file.'),
        ('lark_reader.py', 'Lark Sheet Reader', 'Reads Lark Sheet. Returns rows where Column F has URL and A/D/E/G are empty.'),
        ('lark_writer.py', 'Lark Sheet Writer', 'Writes to columns A/D/E/G ONLY. Hard allowlist raises exception for any other column.'),
        ('browser_reader.py', 'Browser Controller', 'Opens video URLs in persistent Chrome using Playwright. Takes screenshots.'),
        ('vision_extract.py', 'AI Vision Extractor', 'Sends screenshot to Claude Haiku Vision. Returns JSON: posted_date, caption, view_count.'),
        ('processor.py', 'Data Processor', 'Cleans caption (Unicode hashtag strip), formats date, sets Content Type default.'),
        ('reporter.py', 'Run Reporter', 'Prints run summary: rows processed, successes, failures, errors.'),
    ],
    [1.8, 1.8, 4.0]
)

add_heading(doc, '2.2 System Flow Diagram', 2)

for line in [
    '  USER pastes URLs into Lark Column F',
    '         |',
    '  USER runs: python run.py',
    '         |',
    '  [lark_reader.py]',
    '  Reads Lark Sheet -- finds rows where F has URL + A/D/E/G empty',
    '         |',
    '         | for each URL',
    '         v',
    '  [browser_reader.py]',
    '  Opens URL in persistent Chrome (Playwright)',
    '  Takes screenshot of video page',
    '         |',
    '         v',
    '  [vision_extract.py]',
    '  Sends screenshot to Claude Haiku Vision',
    '  Returns JSON: { posted_date, caption, view_count }',
    '         |',
    '         v',
    '  [processor.py]',
    '  clean_caption() -- strips ALL Unicode hashtags',
    '  format_date()   -- standardises date format',
    '  set_content_type() -- returns "Content Casual"',
    '         |',
    '         v',
    '  [lark_writer.py]  <-- HARD ALLOWLIST: A, D, E, G only',
    '  Column A <-- posted_date',
    '  Column D <-- cleaned caption',
    '  Column E <-- "Content Casual"',
    '  Column G <-- view_count',
    '  Column B/C/F/H --> RAISES EXCEPTION immediately',
    '         |',
    '         v',
    '  [reporter.py]',
    '  Prints run summary to terminal',
]:
    add_code(doc, line)

doc.add_paragraph()

# ── 3. TECHNOLOGY STACK ───────────────────────────────────────────────────────
add_heading(doc, '3. Technology Stack', 1)
make_table(doc,
    ['Library / Tool', 'Version', 'Purpose'],
    [
        ('Python', '3.11+', 'Core programming language'),
        ('playwright', 'latest (pinned)', 'Browser automation — opens video pages, takes screenshots'),
        ('anthropic', 'latest (pinned)', 'Claude Haiku Vision API for screenshot analysis + AI recommendations'),
        ('lark-oapi', 'latest (pinned)', 'Official Lark Open API SDK for reading/writing Lark Sheet'),
        ('python-dotenv', 'latest (pinned)', 'Load API keys and secrets from .env file securely'),
        ('pytest', 'latest (pinned)', 'Run automated tests — especially column allowlist tests'),
        ('re (built-in)', 'built-in', 'Unicode-aware regex for hashtag removal'),
        ('pathlib (built-in)', 'built-in', 'File path management for Chrome profile directory'),
        ('json (built-in)', 'built-in', 'Parse Claude Vision response JSON'),
        ('datetime (built-in)', 'built-in', 'Date formatting and standardisation'),
    ],
    [2.2, 1.0, 4.3]
)

# ── 4. DETAILED COMPONENT DESIGN ─────────────────────────────────────────────
add_heading(doc, '4. Detailed Component Design', 1)

add_heading(doc, '4.1 lark_reader.py', 2)
add_body(doc, 'Reads the Lark Sheet and returns rows that need processing.')
make_table(doc,
    ['Function', 'Input', 'Output', 'Logic'],
    [
        ('get_rows_to_process()', 'None', 'List of {row_index, url}', 'Find rows where Column F has URL AND at least one of A/D/E/G is empty'),
        ('connect_lark()', 'None', 'Lark client object', 'Authenticate using App ID + Secret from .env'),
        ('read_sheet()', 'Lark client', 'Raw sheet data', 'Read all rows from configured sheet tab'),
    ],
    [2.2, 1.5, 1.8, 2.0]
)

add_heading(doc, '4.2 lark_writer.py — Hard Column Allowlist', 2)
add_body(doc, 'This is the most critical safety component. The hard allowlist MUST be enforced:', bold=True, color=RED)
doc.add_paragraph()

add_code(doc, '  ALLOWED_COLUMNS = {"A", "D", "E", "G"}  # Hard allowlist')
add_code(doc, '')
add_code(doc, '  def write_cell(row_index, column, value):')
add_code(doc, '      if column not in ALLOWED_COLUMNS:')
add_code(doc, '          raise ValueError(')
add_code(doc, '              f"FORBIDDEN: Cannot write to column {column}. "')
add_code(doc, '              f"Only {ALLOWED_COLUMNS} are allowed."')
add_code(doc, '          )')
add_code(doc, '      # proceed to write to Lark Sheet...')
doc.add_paragraph()

make_table(doc,
    ['Function', 'Description'],
    [
        ('write_cell(row_index, column, value)', 'Writes one cell. Raises ValueError immediately if column not in allowlist.'),
        ('write_row(row_index, data_dict)', 'Writes A/D/E/G for one row. Calls write_cell for each — allowlist enforced per call.'),
        ('verify_write(row_index)', 'Reads back the written row to confirm data was saved correctly.'),
    ],
    [3.0, 4.5]
)

add_heading(doc, '4.3 browser_reader.py', 2)
add_body(doc, 'Opens each video URL in a real Chrome browser and takes a screenshot.')
make_table(doc,
    ['Function', 'Description'],
    [
        ('get_browser()', 'Launches Playwright with persistent Chrome profile at ~/.cache/auto-count/chrome-profile'),
        ('open_url(browser, url)', 'Opens the video URL in a new tab. Waits for page to fully load.'),
        ('take_screenshot(page)', 'Takes a full-page screenshot. Returns image bytes.'),
        ('close_browser(browser)', 'Closes browser cleanly after all URLs are processed.'),
    ],
    [2.5, 5.0]
)

add_body(doc, 'Persistent Chrome profile setup:', italic=True, color=GREY)
add_code(doc, '  PROFILE_PATH = Path.home() / ".cache" / "auto-count" / "chrome-profile"')
add_code(doc, '  PROFILE_PATH.mkdir(parents=True, exist_ok=True)')
add_code(doc, '')
add_code(doc, '  browser = playwright.chromium.launch_persistent_context(')
add_code(doc, '      user_data_dir=str(PROFILE_PATH),')
add_code(doc, '      headless=False,  # visible browser so user can log in if needed')
add_code(doc, '  )')
doc.add_paragraph()

add_heading(doc, '4.4 vision_extract.py', 2)
add_body(doc, 'Sends screenshot to Claude Haiku Vision and returns structured data.')
make_table(doc,
    ['Function', 'Input', 'Output'],
    [
        ('extract_from_screenshot(image_bytes, url)', 'Screenshot bytes + URL', 'JSON: { posted_date, caption, view_count }'),
        ('build_prompt(url)', 'Video URL', 'Prompt string telling Claude what to extract'),
        ('parse_response(response)', 'Claude API response', 'Validated dict with posted_date, caption, view_count'),
        ('validate_extraction(data)', 'Extracted dict', 'Raises error if required fields are missing or invalid'),
    ],
    [3.0, 2.0, 2.5]
)

add_body(doc, 'Claude Vision prompt template:', italic=True, color=GREY)
add_code(doc, '  You are analysing a screenshot of a social media video post.')
add_code(doc, '  Extract exactly these 3 fields and return as JSON only:')
add_code(doc, '  {')
add_code(doc, '    "posted_date": "DD/MM/YYYY format of when the video was posted",')
add_code(doc, '    "caption": "Full video caption text including hashtags",')
add_code(doc, '    "view_count": 123456  (integer, the number of views shown)')
add_code(doc, '  }')
add_code(doc, '  If a field cannot be found, use null.')
add_code(doc, '  Return JSON only. No explanation.')
doc.add_paragraph()

add_heading(doc, '4.5 processor.py', 2)
add_body(doc, 'Cleans and formats extracted data before writing to Lark Sheet.')
make_table(doc,
    ['Function', 'Input', 'Output', 'Rule'],
    [
        ('clean_caption(text)', 'Raw caption with hashtags', 'Caption without any hashtags', 'Unicode-aware: strips #word in any language'),
        ('format_date(raw_date)', 'Date string from Claude', 'DD/MM/YYYY string', 'Standardise to consistent format'),
        ('set_content_type()', 'Nothing', '"Content Casual"', 'Always returns this exact string'),
        ('handle_empty(value, fallback)', 'Any value', 'Value or fallback', 'Returns fallback if value is None or empty'),
    ],
    [2.0, 1.8, 1.8, 1.9]
)

add_body(doc, 'Unicode-aware hashtag removal:', italic=True, color=GREY)
add_code(doc, '  import re')
add_code(doc, '')
add_code(doc, '  def clean_caption(text):')
add_code(doc, '      if not text or text.strip() == "":')
add_code(doc, '          return "No caption"')
add_code(doc, '      # \\w matches Unicode word chars (Chinese, Arabic, emoji, etc.)')
add_code(doc, '      cleaned = re.sub(r\'#\\w+\', \'\', text, flags=re.UNICODE)')
add_code(doc, '      return " ".join(cleaned.split())  # collapse extra whitespace')
doc.add_paragraph()

add_heading(doc, '4.6 run.py', 2)
add_body(doc, 'Main entry point. Orchestrates all components in order.')
add_code(doc, '  1. Load .env (API keys, Lark credentials)')
add_code(doc, '  2. lark_reader.get_rows_to_process() --> list of {row_index, url}')
add_code(doc, '  3. browser_reader.get_browser()')
add_code(doc, '  4. For each row:')
add_code(doc, '       a. browser_reader.open_url(url)')
add_code(doc, '       b. browser_reader.take_screenshot()')
add_code(doc, '       c. vision_extract.extract_from_screenshot()')
add_code(doc, '       d. processor.clean_caption() / format_date() / set_content_type()')
add_code(doc, '       e. lark_writer.write_row(row_index, {A, D, E, G})')
add_code(doc, '       f. Record success or failure')
add_code(doc, '  5. browser_reader.close_browser()')
add_code(doc, '  6. reporter.print_summary(results)')
doc.add_paragraph()

add_heading(doc, '4.7 reporter.py', 2)
add_body(doc, 'Prints a clear run summary to the terminal after all URLs are processed.')
add_code(doc, '  ================================')
add_code(doc, '  AUTO COUNT SOCIAL MEDIA REACH')
add_code(doc, '  Run Summary — 11 May 2026 10:30')
add_code(doc, '  ================================')
add_code(doc, '  Total rows processed : 10')
add_code(doc, '  Successful           : 9')
add_code(doc, '  Failed               : 1')
add_code(doc, '  --------------------------------')
add_code(doc, '  FAILED:')
add_code(doc, '  Row 5 | https://rednote.com/... | Error: login required')
add_code(doc, '  ================================')
doc.add_paragraph()

# ── 5. DATA FLOW ──────────────────────────────────────────────────────────────
add_heading(doc, '5. Data Flow', 1)

make_table(doc,
    ['Step', 'From', 'To', 'Data'],
    [
        ('1', 'User', 'Lark Sheet Column F', 'Video URL (user pastes manually)'),
        ('2', 'lark_reader', 'run.py', 'List of {row_index, url} where A/D/E/G are empty'),
        ('3', 'run.py', 'browser_reader', 'URL string'),
        ('4', 'browser_reader', 'Playwright Chrome', 'Open URL command'),
        ('5', 'Playwright Chrome', 'vision_extract', 'Screenshot image bytes'),
        ('6', 'vision_extract', 'Claude Haiku Vision API', 'Screenshot + extraction prompt'),
        ('7', 'Claude Haiku Vision API', 'vision_extract', 'JSON: { posted_date, caption, view_count }'),
        ('8', 'vision_extract', 'processor', 'Raw extracted data'),
        ('9', 'processor', 'lark_writer', 'Clean data: { A: date, D: caption, E: "Content Casual", G: views }'),
        ('10', 'lark_writer', 'Lark Sheet', 'Writes to columns A, D, E, G in existing row'),
        ('11', 'run.py', 'reporter', 'List of results (success/fail per row)'),
        ('12', 'reporter', 'Terminal', 'Printed run summary'),
    ],
    [0.5, 2.0, 2.0, 3.0]
)

# ── 6. COLUMN ALLOWLIST ENFORCEMENT ──────────────────────────────────────────
add_heading(doc, '6. Column Allowlist Enforcement', 1)
add_body(doc, 'This is the most critical safety mechanism in the system.', bold=True)
doc.add_paragraph()

make_table(doc,
    ['Column', 'Allowed to Write?', 'What Happens if Write Attempted'],
    [
        ('A — Date', ('YES', False, GREEN), 'Write proceeds normally'),
        ('B — Week', ('NO', True, RED), 'ValueError raised immediately — write aborted'),
        ('C — PIC', ('NO', True, RED), 'ValueError raised immediately — write aborted'),
        ('D — Title', ('YES', False, GREEN), 'Write proceeds normally'),
        ('E — Content Type', ('YES', False, GREEN), 'Write proceeds normally — always "Content Casual"'),
        ('F — Link', ('NO', True, RED), 'ValueError raised immediately — user URL protected'),
        ('G — Reach', ('YES', False, GREEN), 'Write proceeds normally'),
        ('H — Final Reach', ('NO', True, RED), 'ValueError raised immediately — write aborted'),
    ],
    [1.5, 1.8, 4.2]
)

add_heading(doc, '6.1 pytest Tests for Allowlist', 2)
add_body(doc, 'These tests MUST pass (green) before development continues:', italic=True, color=GREY)
add_code(doc, '  def test_write_to_B_raises():')
add_code(doc, '      with pytest.raises(ValueError):')
add_code(doc, '          lark_writer.write_cell(1, "B", "value")')
add_code(doc, '')
add_code(doc, '  def test_write_to_C_raises():')
add_code(doc, '      with pytest.raises(ValueError):')
add_code(doc, '          lark_writer.write_cell(1, "C", "value")')
add_code(doc, '')
add_code(doc, '  def test_write_to_F_raises():')
add_code(doc, '      with pytest.raises(ValueError):')
add_code(doc, '          lark_writer.write_cell(1, "F", "value")')
add_code(doc, '')
add_code(doc, '  def test_write_to_H_raises():')
add_code(doc, '      with pytest.raises(ValueError):')
add_code(doc, '          lark_writer.write_cell(1, "H", "value")')
add_code(doc, '')
add_code(doc, '  def test_write_to_A_succeeds():')
add_code(doc, '      # Should NOT raise')
add_code(doc, '      lark_writer.write_cell(1, "A", "01/05/2026")')
doc.add_paragraph()

# ── 7. LARK SHEET INTEGRATION ─────────────────────────────────────────────────
add_heading(doc, '7. Lark Sheet Integration', 1)

add_heading(doc, '7.1 Authentication', 2)
make_table(doc,
    ['Setting', 'Description'],
    [
        ('LARK_APP_ID', 'From Lark Open Platform — stored in .env'),
        ('LARK_APP_SECRET', 'From Lark Open Platform — stored in .env'),
        ('LARK_SHEET_TOKEN', 'From the Lark Sheet URL — stored in .env'),
        ('LARK_SHEET_TAB', 'The specific tab name where video data lives — stored in .env'),
    ],
    [2.5, 5.0]
)

add_heading(doc, '7.2 Read Logic (lark_reader)', 2)
for item in [
    'Read all rows from the configured sheet tab.',
    'For each row, check if Column F has a non-empty URL.',
    'If Column F has URL AND at least one of A/D/E/G is empty, add to processing list.',
    'Return list of {row_index, url} pairs.',
]:
    add_bullet(doc, item)

add_heading(doc, '7.3 Write Logic (lark_writer)', 2)
for item in [
    'Check column against ALLOWED_COLUMNS allowlist — raise ValueError if not allowed.',
    'Write to the specific row_index — never a new row.',
    'Write columns A, D, E, G in a single batch API call for efficiency.',
    'Read back the row after writing to verify data was saved correctly.',
]:
    add_bullet(doc, item)

# ── 8. CLAUDE VISION INTEGRATION ─────────────────────────────────────────────
add_heading(doc, '8. Claude Vision Integration', 1)

make_table(doc,
    ['Setting', 'Value'],
    [
        ('Model', 'claude-haiku-4-5-20251001'),
        ('API Key', 'ANTHROPIC_API_KEY from .env — never logged'),
        ('Input', 'Base64-encoded screenshot image + extraction prompt'),
        ('Output', 'JSON string: { posted_date, caption, view_count }'),
        ('Fallback', 'If extraction fails, log error and skip row — do not write partial data'),
    ],
    [2.5, 5.0]
)

# ── 9. ERROR HANDLING ─────────────────────────────────────────────────────────
add_heading(doc, '9. Error Handling', 1)
make_table(doc,
    ['Error', 'Where It Occurs', 'System Response'],
    [
        ('Write to forbidden column', 'lark_writer.py', 'Raise ValueError immediately. Log error. Skip this row. Continue.'),
        ('Platform blocks browser', 'browser_reader.py', 'Log error. Skip this URL. Continue with next.'),
        ('Claude Vision returns null fields', 'vision_extract.py', 'Log warning. Use "No caption" / 0 fallbacks. Continue.'),
        ('Lark API auth failure', 'lark_reader/writer', 'Retry token refresh once. If still fails, stop and alert user.'),
        ('URL not recognised as supported platform', 'run.py', 'Log warning. Skip this URL. Continue.'),
        ('Screenshot is blank or loading failed', 'browser_reader.py', 'Retry once after 3 seconds. If still blank, skip row.'),
    ],
    [2.2, 2.0, 3.3]
)

# ── 10. SECURITY ──────────────────────────────────────────────────────────────
add_heading(doc, '10. Security Design', 1)
make_table(doc,
    ['Security Concern', 'How It Is Handled'],
    [
        ('API keys', 'Stored in .env only. Never hardcoded. Never logged. Never printed.'),
        ('.env file', 'Added to .gitignore — never pushed to GitHub.'),
        ('.env.template', 'A safe template with placeholder values pushed to GitHub for reference.'),
        ('Lark credentials', 'App ID and Secret in .env. Token refreshed per session.'),
        ('Chrome profile', 'Stored locally at ~/.cache/auto-count/chrome-profile. Never shared.'),
        ('Column protection', 'Hard allowlist in code raises exception — cannot be bypassed at runtime.'),
    ],
    [2.5, 5.0]
)

# ── 11. PROJECT FILE STRUCTURE ────────────────────────────────────────────────
add_heading(doc, '11. Project File Structure', 1)
for line in [
    'Auto-Count-Social-MediaReach/',
    '|-- run.py                    # Main entry point: python run.py',
    '|-- lark_reader.py            # Read Lark Sheet, find rows to process',
    '|-- lark_writer.py            # Write to Lark Sheet (hard allowlist A/D/E/G)',
    '|-- browser_reader.py         # Playwright Chrome automation + screenshots',
    '|-- vision_extract.py         # Claude Haiku Vision: screenshot -> JSON',
    '|-- processor.py              # clean_caption, format_date, set_content_type',
    '|-- reporter.py               # Print run summary to terminal',
    '|-- tests/',
    '|   |-- test_lark_writer.py   # pytest: B/C/F/H writes raise ValueError',
    '|   |-- test_processor.py     # pytest: hashtag removal tests',
    '|-- .env                      # API keys (NEVER pushed to GitHub)',
    '|-- .env.template             # Safe template with placeholder values',
    '|-- .gitignore                # Excludes .env, chrome-profile, __pycache__',
    '|-- requirements.txt          # Pinned Python dependencies',
    '|-- BRD_Auto_Count_Social_Media_Reach.docx',
    '|-- TDD_Auto_Count_Social_Media_Reach.docx',
]:
    add_code(doc, '  ' + line)
doc.add_paragraph()

# ── 12. BUILD ORDER ───────────────────────────────────────────────────────────
add_heading(doc, '12. Build Order', 1)
make_table(doc,
    ['Phase', 'What to Build', 'Stop Condition'],
    [
        ('1', 'Project scaffold: requirements.txt (pinned), .gitignore, .env.template, virtualenv', 'Virtual environment activated successfully'),
        ('2', 'lark_writer.py with hard allowlist + pytest tests for B/C/F/H', 'ALL pytest tests GREEN before proceeding'),
        ('3', 'lark_reader.py — find rows where F has URL and A/D/E/G are empty', 'Returns correct rows from test sheet'),
        ('4', 'processor.py — clean_caption, format_date, set_content_type', 'Unit tests pass for hashtag removal'),
        ('5', 'browser_reader.py — Playwright persistent Chrome profile', 'Browser opens URL and returns screenshot'),
        ('6', 'vision_extract.py — Claude Vision screenshot to JSON', 'Returns valid JSON for a YouTube video'),
        ('7', 'run.py + reporter.py — wire everything + run summary', 'Full run completes without errors'),
        ('8', 'End-to-end test on 3 real YouTube URLs', 'Lark Sheet correctly filled for all 3 rows'),
    ],
    [0.5, 4.5, 2.5]
)

# ── 13. TESTING PLAN ──────────────────────────────────────────────────────────
add_heading(doc, '13. Testing Plan', 1)
make_table(doc,
    ['Test', 'Type', 'What Is Tested', 'Pass Criteria'],
    [
        ('test_write_B_raises', 'pytest unit', 'lark_writer rejects Column B write', 'ValueError raised'),
        ('test_write_C_raises', 'pytest unit', 'lark_writer rejects Column C write', 'ValueError raised'),
        ('test_write_F_raises', 'pytest unit', 'lark_writer rejects Column F write', 'ValueError raised'),
        ('test_write_H_raises', 'pytest unit', 'lark_writer rejects Column H write', 'ValueError raised'),
        ('test_write_A_ok', 'pytest unit', 'lark_writer allows Column A write', 'No exception raised'),
        ('test_clean_caption_basic', 'pytest unit', 'Hashtag removal — basic English', 'Zero # chars in output'),
        ('test_clean_caption_unicode', 'pytest unit', 'Hashtag removal — Chinese/Arabic tags', 'Zero # chars in output'),
        ('test_clean_caption_empty', 'pytest unit', 'Empty caption returns "No caption"', 'Returns "No caption"'),
        ('End-to-end YouTube x3', 'Integration', 'Full run on 3 real YouTube URLs', 'Lark Sheet A/D/E/G filled correctly'),
    ],
    [2.0, 1.2, 2.5, 1.8]
)

# ── 14. DEPENDENCIES ──────────────────────────────────────────────────────────
add_heading(doc, '14. requirements.txt (Pinned)', 1)
add_body(doc, 'All versions must be pinned in requirements.txt:')
add_code(doc, '  playwright==1.44.0')
add_code(doc, '  anthropic==0.28.0')
add_code(doc, '  lark-oapi==1.3.5')
add_code(doc, '  python-dotenv==1.0.1')
add_code(doc, '  pytest==8.2.0')
doc.add_paragraph()

# ── 15. SIGN-OFF ──────────────────────────────────────────────────────────────
add_heading(doc, '15. Sign-Off & Approval', 1)
add_body(doc, 'By signing below, the approver confirms that this Technical Design Document accurately represents the proposed technical approach.')
doc.add_paragraph()
make_table(doc,
    ['Role', 'Name', 'Signature', 'Date'],
    [
        ('Prepared By', 'jingyi-rere', '', ''),
        ('Reviewed By', '', '', ''),
        ('Approved By', '', '', ''),
    ],
    [2.0, 2.0, 2.0, 1.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— End of Document — Version 2.0 —')
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(*GREY)

output_path = '/Users/jjjyyy/Auto-Count Social MediaReach/TDD_Auto_Count_Social_Media_Reach.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
