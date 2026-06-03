from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

BLUE   = (31, 73, 125)
LBLUE  = (68, 114, 196)
WHITE  = (255, 255, 255)
GREY   = (89, 89, 89)
GREEN  = (0, 112, 0)
ORANGE = (197, 90, 17)
RED    = (192, 0, 0)

def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.style = f'Heading {level}'
    p.add_run(text)
    return p

def add_body(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color='CCCCCC'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), color)
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def set_col_width(cell, width):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def make_table(doc, headers, rows_data, col_widths, header_color='1F497D', alt_color='F0F4FF'):
    t = doc.add_table(rows=len(rows_data)+1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        set_col_width(cell, col_widths[j])
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(*WHITE)
        shade_cell(cell, header_color)
        set_cell_border(cell, '1F497D')
    for i, row_data in enumerate(rows_data, 1):
        row = t.rows[i]
        fill = alt_color if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            set_col_width(cell, col_widths[j])
            if isinstance(val, tuple):
                text, bold, color = val
                run = cell.paragraphs[0].add_run(str(text))
                run.bold = bold
                run.font.size = Pt(10)
                if color:
                    run.font.color.rgb = RGBColor(*color)
            else:
                run = cell.paragraphs[0].add_run(str(val))
                run.font.size = Pt(10)
            shade_cell(cell, fill)
            set_cell_border(cell)
    doc.add_paragraph()
    return t

# ── COVER PAGE ────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BUSINESS REQUIREMENTS DOCUMENT')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(*BLUE)

doc.add_paragraph()

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Auto Count Social Media Reach')
run2.bold = True
run2.font.size = Pt(18)
run2.font.color.rgb = RGBColor(*LBLUE)

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('AI-Assisted Video Metrics Automation System')
run3.italic = True
run3.font.size = Pt(13)
run3.font.color.rgb = RGBColor(*GREY)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_table(rows=5, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (lbl, val) in enumerate([
    ('Document Version:', '4.0'),
    ('Date:', datetime.date.today().strftime('%d %B %Y')),
    ('Status:', 'Final'),
    ('Prepared By:', 'jingyi-rere'),
    ('Reviewed By:', 'Pending'),
]):
    set_col_width(meta.rows[i].cells[0], 2.2)
    set_col_width(meta.rows[i].cells[1], 4.3)
    lc = meta.rows[i].cells[0].paragraphs[0].add_run(lbl)
    lc.bold = True
    lc.font.size = Pt(10)
    vc = meta.rows[i].cells[1].paragraphs[0].add_run(val)
    vc.font.size = Pt(10)
    shade_cell(meta.rows[i].cells[0], 'D6E4F7')
    set_cell_border(meta.rows[i].cells[0])
    set_cell_border(meta.rows[i].cells[1])

doc.add_paragraph()
doc.add_paragraph()

add_heading(doc, 'Document Version History', 2)
make_table(doc,
    ['Version', 'Date', 'Author', 'Changes'],
    [
        ('1.0', '07 May 2026', 'jingyi-rere', 'Initial draft'),
        ('2.0', '09 May 2026', 'jingyi-rere', 'Added user stories, MoSCoW, As-Is vs To-Be, Risk Register, Roadmap'),
        ('3.0', '09 May 2026', 'jingyi-rere', 'Updated: video content only, Lark Sheet output, exact column mapping A-H'),
        ('4.0', datetime.date.today().strftime('%d %B %Y'), 'jingyi-rere', 'Updated: Playwright + Claude Vision approach (no platform APIs), hard column rules, run.py trigger, write to A/D/E/G only'),
    ],
    [0.7, 1.5, 1.5, 3.8]
)

doc.add_page_break()

# ── 1. EXECUTIVE SUMMARY ──────────────────────────────────────────────────────
add_heading(doc, '1. Executive Summary', 1)
add_body(doc, (
    'This document defines the business requirements for an AI-assisted Social Media Video Metrics '
    'Automation system. The system is for a single video content creator managing video content across '
    '7 social media platforms. The user pastes video URLs directly into their Lark Sheet (Column F). '
    'The system then automatically reads those URLs, opens each video in a real browser using Playwright, '
    'takes a screenshot, and uses Claude AI Vision to extract the posted date, video caption, and view count. '
    'The extracted data is written into the correct Lark Sheet columns automatically. '
    'The system runs via a single command: python run.py. No UI is required.'
))

# ── 2. HARD RULES ─────────────────────────────────────────────────────────────
add_heading(doc, '2. Hard Rules (NEVER Violate)', 1)
add_body(doc, 'These rules are enforced in code and must never be broken under any circumstance:', bold=True, color=RED)
doc.add_paragraph()

make_table(doc,
    ['Rule', 'Description'],
    [
        (('RULE-01', True, RED), 'Only write to Lark columns A, D, E, G. Writing to B, C, F, or H must raise an exception immediately.'),
        (('RULE-02', True, RED), 'Never overwrite Column F. The user pastes URLs into Column F manually — the system must never touch it.'),
        (('RULE-03', True, RED), 'Never append new rows. Always write into the existing row identified by the URL in Column F.'),
        (('RULE-04', True, RED), 'Strip ALL hashtags from caption (Unicode-aware) before writing to Column D. No hashtags may remain.'),
        (('RULE-05', True, RED), 'Column E must always be exactly "Content Casual" — hardcoded, no exceptions.'),
    ],
    [1.2, 6.3]
)

# ── 3. PROBLEM STATEMENT ──────────────────────────────────────────────────────
add_heading(doc, '3. Problem Statement', 1)

add_heading(doc, '3.1 As-Is Process (Current — Manual)', 2)
add_body(doc, 'Every reporting cycle, the user has to manually:')
make_table(doc,
    ['Step', 'Platform', 'What the User Does Manually'],
    [
        ('1', 'Instagram', 'Open each Reel, note posted date, copy caption, record view count'),
        ('2', 'TikTok', 'Open each video, note posted date, copy caption, record view count'),
        ('3', 'Facebook', 'Open each video, note posted date, copy caption, record view count'),
        ('4', 'YouTube', 'Open each video, note posted date, copy title, record view count'),
        ('5', 'X (Twitter)', 'Open each post, note posted date, copy caption, record view count'),
        ('6', 'RedNote', 'Open each video, note posted date, copy caption, record view count'),
        ('7', 'Threads', 'Open each post, note posted date, copy caption, record view count'),
        ('8', 'Lark Sheet', 'Manually type all data into columns A, D, E, G — very error prone'),
    ],
    [0.5, 1.5, 5.5]
)

add_heading(doc, '3.2 To-Be Process (Future — Automated)', 2)
add_body(doc, 'With the new system, the user only does TWO things:')
make_table(doc,
    ['Step', 'What the User Does', 'What the System Does'],
    [
        ('1', 'Paste video URLs into Column F of the Lark Sheet', 'Nothing yet — waits for user to finish'),
        ('2', 'Run: python run.py', 'Reads URLs from Column F, opens each in browser, takes screenshot, uses Claude Vision to extract data, writes to columns A/D/E/G'),
    ],
    [0.5, 2.8, 4.2]
)

add_heading(doc, '3.3 Pain Points', 2)
make_table(doc,
    ['#', 'Pain Point', 'Impact'],
    [
        ('1', 'Opens 7 platforms manually every reporting cycle', 'Very time-consuming'),
        ('2', 'Manually types data into wrong Lark column accidentally', 'Data errors hard to spot'),
        ('3', 'Removes hashtags from captions manually', 'Easy to miss some hashtags'),
        ('4', 'No performance analysis', 'Cannot quickly see which videos worked best'),
        ('5', 'Repetitive task takes time away from content creation', 'Less time for creative work'),
    ],
    [0.5, 3.0, 4.0]
)

# ── 4. BUSINESS OBJECTIVES ────────────────────────────────────────────────────
add_heading(doc, '4. Business Objectives', 1)
for obj in [
    'Save time — reduce reporting time by at least 80% per cycle.',
    'User pastes URLs into Column F — system fills A, D, E, G automatically.',
    'Extract video captions and remove all hashtags (Unicode-aware) automatically.',
    'Default Content Type to "Content Casual" in Column E — always.',
    'Generate AI-powered video performance analysis and content recommendations.',
    'Support all 7 platforms using Playwright browser automation + Claude Vision.',
    'Run with a single command: python run.py. No UI, no manual steps.',
]:
    add_bullet(doc, obj)

# ── 5. USER STORIES ───────────────────────────────────────────────────────────
add_heading(doc, '5. User Stories', 1)
make_table(doc,
    ['ID', 'User Story'],
    [
        ('US-01', 'As a video content creator, I want to paste video URLs into Column F of my Lark Sheet and run python run.py so the system fills columns A, D, E, G automatically.'),
        ('US-02', 'As a user, I want Column A filled with the original video posted date so I do not have to look it up.'),
        ('US-03', 'As a user, I want Column D filled with the video caption with all hashtags removed so my sheet stays clean.'),
        ('US-04', 'As a user, I want Column E always set to "Content Casual" automatically without me selecting it.'),
        ('US-05', 'As a user, I want Column F (my URL) to never be touched or overwritten by the system.'),
        ('US-06', 'As a user, I want Column G filled with the video view count automatically.'),
        ('US-07', 'As a user, I want the system to leave Columns B, C, and H completely untouched always.'),
        ('US-08', 'As a user, I want AI recommendations so I know what type of videos to make next cycle.'),
        ('US-09', 'As a user, I want a clear run summary after python run.py so I know what succeeded and what failed.'),
    ],
    [0.8, 6.7]
)

# ── 6. LARK SHEET COLUMN MAPPING ──────────────────────────────────────────────
add_heading(doc, '6. Lark Sheet Column Mapping', 1)
add_body(doc, 'CRITICAL: The system uses a hard allowlist. Only columns A, D, E, G are writable. Any attempt to write to B, C, F, or H raises an exception.', bold=True, color=RED)
doc.add_paragraph()

SYS  = ('System writes automatically', False, GREEN)
USER = ('User fills manually — NEVER TOUCH', False, RED)
DEF  = ('System fills with hardcoded default', False, LBLUE)

make_table(doc,
    ['Column', 'Field Name', 'Who Fills It', 'Rule / Detail'],
    [
        ('A', 'Date',         SYS,  'Original video posted date extracted via Claude Vision. Format: DD/MM/YYYY'),
        ('B', 'Week',         USER, 'User fills manually. System raises exception if it tries to write here.'),
        ('C', 'PIC',          USER, 'User fills manually. System raises exception if it tries to write here.'),
        ('D', 'Title',        SYS,  'Video caption extracted via Claude Vision. ALL hashtags stripped (Unicode-aware). Write "No caption" if empty.'),
        ('E', 'Content Type', DEF,  'Hardcoded as "Content Casual". Always. No exceptions.No overrides.'),
        ('F', 'Link',         USER, 'USER pastes URL here. System READS this column but NEVER writes to it. Exception raised if write attempted.'),
        ('G', 'Reach',        SYS,  'Video view count extracted via Claude Vision.'),
        ('H', 'Final Reach',  USER, 'User fills manually. System raises exception if it tries to write here.'),
    ],
    [0.6, 1.3, 2.5, 3.1]
)

# ── 7. HOW THE SYSTEM WORKS ───────────────────────────────────────────────────
add_heading(doc, '7. How the System Works', 1)
add_body(doc, 'The system uses NO platform APIs. Instead it uses a real browser + AI vision:', italic=True, color=GREY)
doc.add_paragraph()

make_table(doc,
    ['Step', 'What Happens', 'Technology Used'],
    [
        ('1', 'User pastes video URLs into Column F of Lark Sheet', 'User action'),
        ('2', 'User runs: python run.py', 'Terminal command'),
        ('3', 'System reads Lark Sheet — finds rows where Column F has a URL but A/D/E/G are empty', 'lark_reader.py + Lark API'),
        ('4', 'For each URL, system opens the video in a real Chrome browser', 'Playwright + persistent Chrome profile'),
        ('5', 'System takes a screenshot of the video page', 'Playwright screenshot'),
        ('6', 'Claude AI Vision analyses the screenshot and extracts: posted date, caption, view count', 'claude-haiku-4-5 vision model'),
        ('7', 'System cleans the caption (removes all hashtags)', 'processor.py with Unicode regex'),
        ('8', 'System writes data to columns A, D, E, G only', 'lark_writer.py with hard allowlist'),
        ('9', 'System prints a run summary showing success/failure for each URL', 'reporter.py'),
    ],
    [0.5, 3.5, 3.5]
)

# ── 8. SUPPORTED PLATFORMS ────────────────────────────────────────────────────
add_heading(doc, '8. Supported Platforms', 1)
make_table(doc,
    ['Platform', 'Video Format', 'Column G (Views)', 'Column D (Caption)', 'Data Source'],
    [
        ('Instagram', 'Reels', 'Reels view count', 'Post caption — hashtags removed', 'Playwright + Claude Vision'),
        ('TikTok', 'Video', 'Total video views', 'Video description — hashtags removed', 'Playwright + Claude Vision'),
        ('Facebook', 'Video / Reels', 'Video views', 'Post caption — hashtags removed', 'Playwright + Claude Vision'),
        ('YouTube', 'Video / Shorts', 'Total video views', 'Video title', 'Playwright + Claude Vision'),
        ('X (Twitter)', 'Video post', 'Video views', 'Tweet text — hashtags removed', 'Playwright + Claude Vision'),
        ('RedNote', 'Video', 'Video views', 'Post caption — hashtags removed', 'Playwright + Claude Vision'),
        ('Threads', 'Video post', 'Video views', 'Post caption — hashtags removed', 'Playwright + Claude Vision'),
    ],
    [1.1, 1.1, 1.3, 2.1, 1.9]
)

# ── 9. FUNCTIONAL REQUIREMENTS ────────────────────────────────────────────────
add_heading(doc, '9. Functional Requirements', 1)
add_body(doc, 'Priority: M = Must Have   S = Should Have   C = Could Have', italic=True, color=GREY)
doc.add_paragraph()

M = ('M', False, GREEN)
S = ('S', False, LBLUE)
C = ('C', False, ORANGE)

make_table(doc,
    ['ID', 'Priority', 'Requirement'],
    [
        ('FR-01', M, 'System reads Lark Sheet to find rows where Column F has a URL and columns A/D/E/G are empty.'),
        ('FR-02', M, 'System opens each video URL in a real Chrome browser using Playwright with a persistent profile.'),
        ('FR-03', M, 'System takes a screenshot of the video page.'),
        ('FR-04', M, 'System sends screenshot to Claude Haiku Vision model to extract posted date, caption, and view count.'),
        ('FR-05', M, 'System strips ALL hashtags from caption using Unicode-aware regex before writing to Column D.'),
        ('FR-06', M, 'System writes posted date to Column A only.'),
        ('FR-07', M, 'System writes cleaned caption to Column D only. Writes "No caption" if empty.'),
        ('FR-08', M, 'System writes "Content Casual" to Column E only. Hardcoded always.'),
        ('FR-09', M, 'System writes view count to Column G only.'),
        ('FR-10', M, 'System NEVER writes to Columns B, C, F, or H. Any attempt raises an exception immediately.'),
        ('FR-11', M, 'System writes into existing rows only — never appends new rows.'),
        ('FR-12', M, 'System triggers via python run.py only. No UI required.'),
        ('FR-13', M, 'System supports all 7 platforms: Instagram, TikTok, Facebook, YouTube, X, RedNote, Threads.'),
        ('FR-14', S, 'System prints a run summary showing success/failure status for each URL processed.'),
        ('FR-15', S, 'System generates AI performance analysis and content recommendations after processing.'),
        ('FR-16', S, 'If a URL fails, system logs the error and continues processing remaining URLs.'),
        ('FR-17', C, 'System compares current cycle performance against previous cycle.'),
    ],
    [0.8, 0.8, 5.9]
)

# ── 10. NON-FUNCTIONAL REQUIREMENTS ──────────────────────────────────────────
add_heading(doc, '10. Non-Functional Requirements', 1)
make_table(doc,
    ['ID', 'Category', 'Requirement'],
    [
        ('NFR-01', 'Performance', 'Process up to 20 video URLs within 5 minutes.'),
        ('NFR-02', 'Accuracy', 'View counts and captions must match what is shown on the video page with 99% accuracy.'),
        ('NFR-03', 'Data Integrity', 'System must never write to Columns B, C, F, or H under any circumstance.'),
        ('NFR-04', 'Data Integrity', 'System must never overwrite existing data in columns A, D, E, G if already filled.'),
        ('NFR-05', 'Reliability', 'If one URL fails, system must continue processing all remaining URLs.'),
        ('NFR-06', 'Security', 'API keys stored in .env file only. Never logged or printed. .env excluded from git.'),
        ('NFR-07', 'Testability', 'lark_writer.py must have pytest tests proving B/C/F/H writes raise exceptions.'),
        ('NFR-08', 'Compatibility', 'System runs on Python 3.11+ on macOS.'),
    ],
    [0.9, 1.6, 5.0]
)

# ── 11. RISK REGISTER ─────────────────────────────────────────────────────────
add_heading(doc, '11. Risk Register', 1)
H2 = ('High', False, RED)
M2 = ('Medium', False, ORANGE)
L2 = ('Low', False, GREEN)
make_table(doc,
    ['ID', 'Risk', 'Likelihood', 'Impact', 'Mitigation'],
    [
        ('R-01', 'Platform blocks Playwright browser (detects automation)', M2, H2, 'Use persistent Chrome profile with real login session; use slow human-like interactions'),
        ('R-02', 'Claude Vision extracts wrong data from screenshot', L2, H2, 'Validate extracted data format; flag outliers for manual review'),
        ('R-03', 'System writes to wrong Lark column', L2, H2, 'Hard allowlist in lark_writer.py raises exception immediately if wrong column attempted'),
        ('R-04', 'Hashtag regex misses Unicode hashtags (e.g. Chinese, Arabic)', M2, M2, 'Use Unicode-aware regex pattern covering all Unicode word characters'),
        ('R-05', 'RedNote requires login — Playwright session expires', H2, M2, 'Persistent Chrome profile retains login session; user re-logs if expired'),
        ('R-06', 'Lark Sheet API token expires during run', L2, M2, 'Implement token refresh logic; retry on auth failure'),
    ],
    [0.6, 2.3, 1.0, 0.8, 2.8]
)

# ── 12. ASSUMPTIONS & CONSTRAINTS ────────────────────────────────────────────
add_heading(doc, '12. Assumptions & Constraints', 1)

add_heading(doc, '12.1 Assumptions', 2)
for item in [
    'The user pastes video URLs into Column F of the Lark Sheet before running the system.',
    'The user is already logged into all 7 platforms in the persistent Chrome profile.',
    'The Lark Sheet exists with exactly columns A to H as defined in Section 6.',
    'Videos will always be publicly visible when the browser opens the URL.',
    'The user runs the system on macOS with Python 3.11+ installed.',
    '"Content Casual" is the correct default for all videos in Phase 1.',
]:
    add_bullet(doc, item)

add_heading(doc, '12.2 Constraints', 2)
for item in [
    'System MUST NEVER write to Columns B, C, F, or H — enforced with hard exception in code.',
    'Column F is owned by the user — the system only reads it, never writes to it.',
    'System writes into existing rows only — never creates new rows.',
    'No platform APIs are used — all data is extracted via Playwright + Claude Vision.',
    'Single user only in Phase 1.',
    'System runs via python run.py only — no web UI, no mobile app.',
]:
    add_bullet(doc, item)

# ── 13. IMPLEMENTATION ROADMAP ───────────────────────────────────────────────
add_heading(doc, '13. Implementation Roadmap', 1)
make_table(doc,
    ['Phase', 'What Gets Built', 'Key Output'],
    [
        ('Phase 1\nScaffold', 'Project setup: requirements.txt, .gitignore, .env.template, virtualenv', 'Clean project structure ready to build'),
        ('Phase 2\nLark Writer', 'lark_writer.py with hard column allowlist + pytest tests proving B/C/F/H raise exceptions', 'Green pytest output — safety guaranteed'),
        ('Phase 3\nLark Reader', 'lark_reader.py — find rows where F has URL and A/D/E/G are empty', 'System knows what to process'),
        ('Phase 4\nProcessor', 'processor.py — clean_caption (Unicode hashtag strip), format_date, set_content_type', 'Clean data ready to write'),
        ('Phase 5\nBrowser', 'browser_reader.py — Playwright with persistent Chrome profile', 'Browser opens video pages automatically'),
        ('Phase 6\nVision', 'vision_extract.py — screenshot to Claude Haiku Vision to JSON', 'AI extracts posted_date, caption, view_count'),
        ('Phase 7\nWiring', 'run.py + reporter.py — wire everything together + run summary', 'Single command runs full system'),
        ('Phase 8\nTesting', 'End-to-end test on 3 real YouTube URLs', 'Proven working system'),
    ],
    [1.2, 4.3, 2.0]
)

# ── 14. SUCCESS CRITERIA ─────────────────────────────────────────────────────
add_heading(doc, '14. Success Criteria', 1)
make_table(doc,
    ['Criteria', 'Target'],
    [
        ('Time saved per reporting cycle', 'At least 80% reduction'),
        ('Columns B, C, F, H protection', '100% — exception raised if write attempted'),
        ('Column F (user URL)', 'Never overwritten — always preserved'),
        ('View count accuracy', '99% match with what is shown on the video page'),
        ('Caption accuracy', 'Correct caption with zero hashtags remaining'),
        ('Platforms supported', 'All 7: Instagram, TikTok, Facebook, YouTube, X, RedNote, Threads'),
        ('pytest tests', 'All green — B/C/F/H write tests must pass'),
        ('Trigger method', 'python run.py — single command, no UI needed'),
    ],
    [3.5, 4.0]
)

# ── 15. GLOSSARY ─────────────────────────────────────────────────────────────
add_heading(doc, '15. Glossary', 1)
make_table(doc,
    ['Term', 'Definition'],
    [
        ('Playwright', 'A Python library that controls a real Chrome browser automatically.'),
        ('Claude Vision', 'Claude AI\'s ability to look at a screenshot and extract information from it.'),
        ('Persistent Chrome Profile', 'A saved browser session that remembers your logins across runs.'),
        ('Hard Allowlist', 'A strict list of allowed columns in code — any other column raises an exception immediately.'),
        ('Unicode-aware Regex', 'A pattern that matches hashtags in any language, including Chinese, Arabic, and emoji-based tags.'),
        ('run.py', 'The single Python file the user runs to trigger the entire system.'),
        ('lark_writer.py', 'The module that writes data to Lark Sheet — enforces the column allowlist.'),
        ('lark_reader.py', 'The module that reads the Lark Sheet to find rows needing processing.'),
        ('vision_extract.py', 'The module that sends screenshots to Claude Vision and returns structured data.'),
        ('Content Casual', 'The hardcoded default value always written to Column E.'),
        ('Final Reach', 'Column H — manually calculated by the user. System never touches this.'),
    ],
    [2.0, 5.5]
)

doc.add_paragraph()

# ── 16. SIGN-OFF ──────────────────────────────────────────────────────────────
add_heading(doc, '16. Sign-Off & Approval', 1)
add_body(doc, 'By signing below, the approver confirms that this BRD accurately represents the requirements for the Auto Count Social Media Reach system.')
doc.add_paragraph()
make_table(doc,
    ['Role', 'Name', 'Signature', 'Date'],
    [
        ('Prepared By', 'jingyi-rere', '', ''),
        ('Reviewed By', '', '', ''),
        ('Approved By', '', '', ''),
    ],
    [2.0, 2.0, 2.0, 1.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— End of Document — Version 4.0 —')
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(*GREY)

output_path = '/Users/jjjyyy/Auto-Count Social MediaReach/BRD_Auto_Count_Social_Media_Reach.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
