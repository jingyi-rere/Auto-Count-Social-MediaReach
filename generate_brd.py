from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

BLUE   = (31, 73, 125)
LBLUE  = (68, 114, 196)
WHITE  = (255, 255, 255)
GREY   = (89, 89, 89)
GREEN  = (0, 112, 0)
ORANGE = (197, 90, 17)
RED    = (192, 0, 0)

def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.style = f'Heading {level}'
    p.add_run(text)
    return p

def add_body(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color='CCCCCC'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), color)
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def set_col_width(cell, width):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def make_table(doc, headers, rows_data, col_widths, header_color='1F497D', alt_color='F0F4FF'):
    t = doc.add_table(rows=len(rows_data)+1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        set_col_width(cell, col_widths[j])
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(*WHITE)
        shade_cell(cell, header_color)
        set_cell_border(cell, '1F497D')
    for i, row_data in enumerate(rows_data, 1):
        row = t.rows[i]
        fill = alt_color if i % 2 == 0 else 'FFFFFF'
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            set_col_width(cell, col_widths[j])
            if isinstance(val, tuple):
                text, bold, color = val
                run = cell.paragraphs[0].add_run(str(text))
                run.bold = bold
                run.font.size = Pt(10)
                if color:
                    run.font.color.rgb = RGBColor(*color)
            else:
                run = cell.paragraphs[0].add_run(str(val))
                run.font.size = Pt(10)
            shade_cell(cell, fill)
            set_cell_border(cell)
    doc.add_paragraph()
    return t

# ── COVER PAGE ────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BUSINESS REQUIREMENTS DOCUMENT')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(*BLUE)

doc.add_paragraph()

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Auto Count Social Media Reach')
run2.bold = True
run2.font.size = Pt(18)
run2.font.color.rgb = RGBColor(*LBLUE)

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('AI-Assisted Video Metrics Automation System')
run3.italic = True
run3.font.size = Pt(13)
run3.font.color.rgb = RGBColor(*GREY)

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_table(rows=5, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (lbl, val) in enumerate([
    ('Document Version:', '5.5'),
    ('Date:', datetime.date.today().strftime('%d %B %Y')),
    ('Status:', 'Final'),
    ('Prepared By:', 'jingyi-rere'),
    ('Reviewed By:', 'Pending'),
]):
    set_col_width(meta.rows[i].cells[0], 2.2)
    set_col_width(meta.rows[i].cells[1], 4.3)
    lc = meta.rows[i].cells[0].paragraphs[0].add_run(lbl)
    lc.bold = True
    lc.font.size = Pt(10)
    vc = meta.rows[i].cells[1].paragraphs[0].add_run(val)
    vc.font.size = Pt(10)
    shade_cell(meta.rows[i].cells[0], 'D6E4F7')
    set_cell_border(meta.rows[i].cells[0])
    set_cell_border(meta.rows[i].cells[1])

doc.add_paragraph()
doc.add_paragraph()

add_heading(doc, 'Document Version History', 2)
make_table(doc,
    ['Version', 'Date', 'Author', 'Changes'],
    [
        ('1.0', '07 May 2026', 'jingyi-rere', 'Initial draft'),
        ('2.0', '09 May 2026', 'jingyi-rere', 'Added user stories, MoSCoW, As-Is vs To-Be, Risk Register, Roadmap'),
        ('3.0', '09 May 2026', 'jingyi-rere', 'Updated: video content only, Lark Sheet output, exact column mapping A-H'),
        ('4.0', '11 May 2026', 'jingyi-rere', 'Updated: Playwright + Claude Vision approach (no platform APIs), hard column rules, run.py trigger, write to A/D/E/G only'),
        ('5.0', '13 May 2026', 'jingyi-rere', 'Updated: watcher.py background daemon (every 5 min), yt-dlp for YouTube/TikTok, Firefox for Instagram/RedNote, 283 automated tests, structured logging, retry logic, startup validation, security hardening'),
        ('5.1', '13 May 2026', 'jingyi-rere', 'Added boss-required sections: specific pain with time measurements, why AI not just automation, quantified value (33 hrs/year), AC-01 to AC-10 acceptance criteria'),
        ('5.2', '13 May 2026', 'jingyi-rere', 'Feedback fixes: clarified Phase 1 vs Phase 2 platform scope, reconciled baseline metric (40 min/week), aligned AC-03 and NFR-02 accuracy thresholds'),
        ('5.3', '18 May 2026', 'jingyi-rere', 'Instagram view count fix: now extracted from reel page screenshot (taken before Escape key) — guarantees correct reel. Watcher interval reduced to 3 minutes. Grid navigation retained for caption/date only.'),
        ('5.4', '18 May 2026', 'jingyi-rere', 'Feedback fixes: removed defensive baseline note — stated upfront as 4 min/video × 10 videos. Phase 2 platforms moved to Appendix A (Future Scope) to keep main doc focused on current build.'),
        ('5.5', datetime.date.today().strftime('%d %B %Y'), 'jingyi-rere', 'TikTok → Firefox+DOM (not yt-dlp); X (Twitter) added as 5th platform; live 5-platform test confirmed.'),
    ],
    [0.7, 1.5, 1.5, 3.8]
)

doc.add_page_break()

# ── 1. EXECUTIVE SUMMARY ──────────────────────────────────────────────────────
add_heading(doc, '1. Executive Summary', 1)
add_body(doc, (
    'A video content creator at ricebowlmy spends approximately 40 minutes every week manually '
    'collecting performance data (posted date, caption, view count) from 4 social media platforms '
    '(Instagram, YouTube, TikTok, RedNote) and typing it into a Lark Bitable tracking sheet. '
    'This is pure manual data entry — repetitive, error-prone, and done at the end of an already busy week. '
    '\n\n'
    'This system eliminates that 40 minutes. The user pastes video URLs into Lark Column F. '
    'A background watcher detects new URLs within 3 minutes and fills in all data automatically: '
    'posted date (Column A), caption with hashtags removed (Column D), content type (Column E), '
    'and view count (Column G). '
    '\n\n'
    'Claude AI Vision is required — not optional — because three of the four platforms (Instagram, '
    'TikTok, RedNote) have no public API available, block automated scrapers, and present data '
    'visually in ways that change with every UI update. Only AI vision can reliably read any '
    'screenshot regardless of platform, language, or UI version. '
    '\n\n'
    'Projected outcome: 38 minutes saved per week, approximately 33 hours per year, '
    'with near-zero data entry errors.'
))

# ── 2. HARD RULES ─────────────────────────────────────────────────────────────
add_heading(doc, '2. Hard Rules (NEVER Violate)', 1)
add_body(doc, 'These rules are enforced in code and must never be broken under any circumstance:', bold=True, color=RED)
doc.add_paragraph()

make_table(doc,
    ['Rule', 'Description'],
    [
        (('RULE-01', True, RED), 'Only write to Lark columns A, D, E, G. Writing to B, C, F, or H must raise an exception immediately.'),
        (('RULE-02', True, RED), 'Never overwrite Column F. The user pastes URLs into Column F manually — the system must never touch it.'),
        (('RULE-03', True, RED), 'Never append new rows. Always write into the existing row identified by the URL in Column F.'),
        (('RULE-04', True, RED), 'Strip ALL hashtags from caption (Unicode-aware) before writing to Column D. No hashtags may remain.'),
        (('RULE-05', True, RED), 'Column E must always be exactly "Content Casual" — hardcoded, no exceptions.'),
    ],
    [1.2, 6.3]
)

# ── 3. PROBLEM STATEMENT ──────────────────────────────────────────────────────
add_heading(doc, '3. Problem Statement', 1)

add_heading(doc, '3.1 The Real Daily-Work Pain', 2)
add_body(doc, (
    'Every week, the content creator must manually collect performance metrics for all videos '
    'posted that week across 4 active platforms: Instagram, YouTube, TikTok, and RedNote. '
    'The process is done manually, platform by platform, video by video. '
    'Below is the measured time breakdown for a typical week of 10 videos:'
))
doc.add_paragraph()

make_table(doc,
    ['Task (per video)', 'Time Taken', 'Specific Problem'],
    [
        ('Open platform, find the video', '~1 min', 'Instagram grid is paginated. RedNote UI is in Chinese. Older posts require scrolling.'),
        ('Read and copy the caption', '~1 min', 'Captions are long. Must select carefully to avoid copying overlaid text or comments.'),
        ('Remove hashtags manually', '~1 min', 'A post has 15-25 hashtags in English, Malay, and Chinese. Must delete each one individually. Easy to miss hidden tags.'),
        ('Note the view count', '~30 sec', 'Instagram desktop hides view counts. Must navigate to profile grid to find the number. RedNote shows it as an icon overlay.'),
        ('Type data into Lark Bitable', '~30 sec', 'Switch app, find the right row, type in 4 fields. Typos happen — e.g. 10,000 entered as 1,000.'),
        ('TOTAL per video', '~4 minutes', ''),
        ('TOTAL per week (10 videos)', '~40 minutes', '40 minutes of pure manual data entry every single week.'),
    ],
    [2.5, 1.2, 3.8]
)
doc.add_paragraph()

add_body(doc, (
    'This is not a complex task — it is a repetitive, low-value task that happens every week '
    'without exception. It takes 40 minutes of focused attention that could otherwise go into '
    'content planning, editing, or strategy. Over a year, this is approximately 33 hours lost '
    'to data entry.'
), italic=True, color=GREY)

add_body(doc, (
    'Baseline: 4 minutes per video × 10 videos = 40 minutes per week. '
    'Each task above was timed individually — this is the repeatable, verifiable figure used for all calculations in this document.'
), italic=True, color=GREY)

add_heading(doc, '3.2 Why This Cannot Be Solved by Simple Automation', 2)
add_body(doc, (
    'The obvious approach — write a web scraper — does not work. Here is why each platform resists it:'
))
doc.add_paragraph()

make_table(doc,
    ['Platform', 'Why Simple Automation Fails'],
    [
        ('Instagram', 'Actively blocks scrapers with 429/403 errors and CAPTCHA. HTML structure changes with every UI update. View counts are NOT in the page HTML — they are rendered inside JavaScript components and only visible after login.'),
        ('RedNote (\u5c0f\u7ea2\u4e66)', 'Has NO public API. Chinese-language UI requires understanding of layout context, not just HTML selectors. View count is shown as an icon overlay on the video thumbnail — not in any text element.'),
        ('TikTok', 'API requires business account approval. Scraping is blocked. Numbers like "1.2M" require conversion — a scraper returns the string, not the integer.'),
        ('YouTube', 'API has strict daily quota limits (10,000 units/day). yt-dlp is the only reliable free method, but it still requires intelligent parsing of metadata.'),
    ],
    [1.5, 6.0]
)
doc.add_paragraph()

add_heading(doc, '3.3 As-Is vs To-Be Process', 2)
make_table(doc,
    ['', 'As-Is (Manual)', 'To-Be (Automated)'],
    [
        ('User action required', 'Open each platform, find video, copy data, remove hashtags, type into Lark — for every video', 'Paste URL into Lark Column F. Done. System handles everything else.'),
        ('Time per week (10 videos)', '~40 minutes active work', '~2 minutes (just pasting URLs)'),
        ('Time saved per week', '—', '~38 minutes'),
        ('Time saved per year', '—', '~33 hours'),
        ('Hashtag deletions per week', '~150-200 manual deletions (15-20 tags x 10 videos)', '0 — system removes all automatically'),
        ('Data entry errors', 'Estimated 5-10% error rate (typos, missed hashtags, wrong row)', 'Target <2% (AI reads directly from screen)'),
        ('Works when you are busy?', 'No — requires full attention', 'Yes — runs in background while you work'),
    ],
    [2.2, 2.6, 2.7]
)

# ── 4. WHY AI IS THE RIGHT TOOL ──────────────────────────────────────────────
add_heading(doc, '4. Why AI — Not Just Any Automation', 1)
add_body(doc, (
    'This problem specifically requires AI vision. A rule-based scraper or API integration '
    'cannot solve it. Here is the direct comparison:'
))
doc.add_paragraph()

make_table(doc,
    ['Approach', 'Why It Does NOT Work for This Problem'],
    [
        ('Web scraper (Selenium / BeautifulSoup)',
         'Instagram and TikTok actively detect and block scrapers. '
         'HTML structure changes every few weeks with UI updates — the scraper breaks and requires manual fixing. '
         'View counts on Instagram are rendered in JavaScript after login; they do not exist in raw HTML.'),
        ('Platform APIs (official)',
         'RedNote has NO public API — completely unavailable. '
         'Instagram Graph API requires Facebook Business account verification (weeks of approval process). '
         'YouTube Data API has a 10,000 unit/day quota — enough for ~100 videos, not scalable. '
         'TikTok API requires business account approval.'),
        ('Zapier / Make (no-code automation)',
         'Cannot access platforms without APIs. '
         'Cannot read data from screenshots or handle visual content. '
         'Cannot perform intelligent text operations like "remove hashtags in any language".'),
        ('Rule-based screenshot parser (OCR)',
         'OCR reads raw text but cannot understand context. '
         'Cannot distinguish a view count from a like count, a comment count, or a share count — '
         'all are numbers near icons on screen. '
         'Cannot handle "1.2K" or "3.4M" shorthand conversion reliably across languages.'),
    ],
    [2.2, 5.3]
)
doc.add_paragraph()

add_body(doc, 'Why Claude AI Vision specifically solves what others cannot:', bold=True)
doc.add_paragraph()

make_table(doc,
    ['Capability Required', 'Why AI Vision Is Needed'],
    [
        ('Read any platform layout',
         'Claude Vision reads a screenshot like a human reads a screen. '
         'It does not depend on specific HTML elements or CSS selectors. '
         'When Instagram or RedNote updates their UI, the system still works.'),
        ('Understand visual context',
         'The view count is a number near a play icon. The like count is a number near a heart. '
         'Claude Vision understands the meaning of icons and their associated numbers — '
         'a rule-based system cannot.'),
        ('Handle multi-language content',
         'Captions mix English, Malay, Chinese, and Arabic. Hashtags appear in all these languages. '
         'Claude Vision understands all of them. A regex alone would miss Chinese hashtags (#\u89c6\u9891).'),
        ('Interpret relative dates',
         '"3 days ago", "1w", "2\u9031\u524d" (Chinese for 2 weeks ago) — these require language understanding, '
         'not pattern matching. Claude converts them correctly.'),
        ('Convert number shorthand',
         '"1.2K views", "3.4M plays", "10.3\u4e07" (Chinese for 103,000) — '
         'Claude converts all formats to exact integers. A scraper would return the raw string.'),
    ],
    [2.2, 5.3]
)

# ── 5. QUANTIFIED VALUE ───────────────────────────────────────────────────────
add_heading(doc, '5. Quantified Value', 1)
add_body(doc, 'Based on measured manual workflow times for a typical week of 10 videos:')
doc.add_paragraph()

make_table(doc,
    ['Metric', 'Before (Manual)', 'After (Automated)', 'Improvement'],
    [
        ('Time spent per week on data collection', '~40 minutes', '~2 minutes (paste URLs only)', '95% reduction'),
        ('Time spent per year', '~33 hours', '~1.5 hours', '31.5 hours saved per year'),
        ('Hashtag deletions per week', '~175 (avg 17.5 tags x 10 videos)', '0', '100% eliminated'),
        ('Manual data entry keystrokes per week', '~400 keystrokes (4 fields x 10 videos)', '0', '100% eliminated'),
        ('Data entry error rate', 'Est. 5-10% (typos, missed tags, wrong row)', 'Target <2%', '60-80% fewer errors'),
        ('Platforms requiring manual login to check', '4 platforms opened separately', '0 manual logins needed', '100% eliminated'),
        ('Time from video posted to data in Lark', 'End of week batch (up to 7 days delay)', 'Within 3 minutes of URL paste', 'Near real-time'),
    ],
    [3.0, 1.8, 1.8, 1.4]
)
doc.add_paragraph()
add_body(doc,
    'At a conservative estimate of 38 minutes saved per week, the system pays for its development '
    'cost within the first month of use.',
    italic=True, color=GREY
)

# ── 6. BUSINESS OBJECTIVES ────────────────────────────────────────────────────
add_heading(doc, '6. Business Objectives', 1)
for obj in [
    'Reduce weekly data collection time from ~40 minutes to ~2 minutes (95% reduction).',
    'Eliminate all manual hashtag deletion — system strips hashtags in any language automatically.',
    'Eliminate all manual data entry into Lark — system writes A, D, E, G with zero keystrokes from user.',
    'Achieve data accuracy of 98%+ — AI reads directly from screen, no transcription errors.',
    'Support all 4 active platforms: Instagram, YouTube, TikTok, RedNote.',
    'Process new URLs within 3 minutes of being pasted — no batch end-of-week delay.',
    'Run unattended — watcher.py runs in background, user does not need to trigger it per URL.',
]:
    add_bullet(doc, obj)

# ── 5. USER STORIES ───────────────────────────────────────────────────────────
add_heading(doc, '7. User Stories', 1)
make_table(doc,
    ['ID', 'User Story'],
    [
        ('US-01', 'As a video content creator, I want to paste video URLs into Column F and have the system fill columns A, D, E, G automatically within 3 minutes — so I save the 40 minutes I currently spend on manual data entry every week.'),
        ('US-02', 'As a user, I want Column A filled with the original video posted date so I do not have to look it up.'),
        ('US-03', 'As a user, I want Column D filled with the video caption with all hashtags removed so my sheet stays clean.'),
        ('US-04', 'As a user, I want Column E always set to "Content Casual" automatically without me selecting it.'),
        ('US-05', 'As a user, I want Column F (my URL) to never be touched or overwritten by the system.'),
        ('US-06', 'As a user, I want Column G filled with the video view count automatically.'),
        ('US-07', 'As a user, I want the system to leave Columns B, C, and H completely untouched always.'),
        ('US-08', 'As a user, I want the system to work even when platforms update their UI — so I do not spend time fixing a broken scraper.'),
        ('US-09', 'As a user, I want error messages in plain English so I know exactly what went wrong and what to do, without needing IT knowledge.'),
    ],
    [0.8, 6.7]
)

# ── 6. LARK SHEET COLUMN MAPPING ──────────────────────────────────────────────
add_heading(doc, '8. Lark Sheet Column Mapping', 1)
add_body(doc, 'CRITICAL: The system uses a hard allowlist. Only columns A, D, E, G are writable. Any attempt to write to B, C, F, or H raises an exception.', bold=True, color=RED)
doc.add_paragraph()

SYS  = ('System writes automatically', False, GREEN)
USER = ('User fills manually — NEVER TOUCH', False, RED)
DEF  = ('System fills with hardcoded default', False, LBLUE)

make_table(doc,
    ['Column', 'Field Name', 'Who Fills It', 'Rule / Detail'],
    [
        ('A', 'Date',         SYS,  'Original video posted date extracted via Claude Vision. Format: DD/MM/YYYY'),
        ('B', 'Week',         USER, 'User fills manually. System raises exception if it tries to write here.'),
        ('C', 'PIC',          USER, 'User fills manually. System raises exception if it tries to write here.'),
        ('D', 'Title',        SYS,  'Video caption extracted via Claude Vision. ALL hashtags stripped (Unicode-aware). Write "No caption" if empty.'),
        ('E', 'Content Type', DEF,  'Hardcoded as "Content Casual". Always. No exceptions.No overrides.'),
        ('F', 'Link',         USER, 'USER pastes URL here. System READS this column but NEVER writes to it. Exception raised if write attempted.'),
        ('G', 'Reach',        SYS,  'Video view count extracted via Claude Vision.'),
        ('H', 'Final Reach',  USER, 'User fills manually. System raises exception if it tries to write here.'),
    ],
    [0.6, 1.3, 2.5, 3.1]
)

# ── 7. HOW THE SYSTEM WORKS ───────────────────────────────────────────────────
add_heading(doc, '9. How the System Works', 1)
add_body(doc, 'The system uses NO platform APIs. Instead it uses a real browser + AI vision:', italic=True, color=GREY)
doc.add_paragraph()

make_table(doc,
    ['Step', 'What Happens', 'Technology Used'],
    [
        ('1', 'User starts watcher once: python watcher.py', 'Terminal command'),
        ('2', 'User pastes video URLs into Column F of Lark Bitable', 'User action'),
        ('3', 'Watcher checks Lark every 3 minutes — finds rows where Column F has URL but A/D/E/G are empty', 'watcher.py + lark_reader.py + Lark API'),
        ('4a', 'YouTube / TikTok URLs: metadata extracted directly (no browser needed)', 'yt-dlp + Chrome cookies'),
        ('4b', 'Instagram / RedNote URLs: video opened in Firefox, screenshot taken', 'Playwright + persistent Firefox profile'),
        ('4c', 'Instagram: two screenshots — reel page (caption/date) + grid page (view count)', 'Firefox + Claude Vision (claude-haiku-4-5)'),
        ('5', 'System cleans the caption (removes all hashtags, Unicode-aware)', 'src/utils.py clean_caption()'),
        ('6', 'System writes all data to columns A, D, E, G in ONE atomic Lark API call', 'lark_writer.py with hard allowlist + retry'),
        ('7', 'All activity logged to logs/auto_count.log. Errors shown in plain English.', 'src/logger.py + src/friendly_errors.py'),
    ],
    [0.5, 3.5, 3.5]
)

# ── 8. SUPPORTED PLATFORMS ────────────────────────────────────────────────────
add_heading(doc, '10. Supported Platforms (Current Scope)', 1)
add_body(doc, (
    'The following 5 platforms are fully built, tested, and active. '
    'They cover 100% of current posting activity.'
), italic=True, color=GREY)
doc.add_paragraph()

make_table(doc,
    ['Platform', 'Video Format', 'Column G (Views)', 'Column D (Caption)', 'Data Source'],
    [
        ('YouTube', 'Video / Shorts', 'Total video views', 'Video title', 'yt-dlp (no browser needed — fast & exact)'),
        ('TikTok', 'Video', 'Total video views', 'Video description — hashtags removed', 'Firefox + DOM (caption/date from DOM; view count from embedded page JSON)'),
        ('X (Twitter)', 'Video post', 'Total video views', 'Tweet text — hashtags removed', 'Firefox + DOM (caption from tweetText; date from <time>; view count from post page)'),
        ('Instagram', 'Reels', 'Reels view count', 'Post caption — hashtags removed', 'Firefox + Claude Vision (2-screenshot approach)'),
        ('RedNote (小红书)', 'Video', 'Video views', 'Post caption — hashtags removed', 'Firefox + Claude Vision'),
    ],
    [1.2, 1.0, 1.2, 1.8, 1.5]
)
doc.add_paragraph()

add_heading(doc, 'Appendix A: Future Scope (Not Built)', 1)
add_body(doc, (
    'The following platforms are out of scope for the current build. '
    'They will only be considered if the user expands posting activity to these platforms.'
), italic=True, color=GREY)
doc.add_paragraph()

make_table(doc,
    ['Platform', 'Video Format', 'Planned Data Source', 'Blocker'],
    [
        ('Facebook', 'Video / Reels', 'Firefox + Claude Vision', 'Requires Facebook Business account for API; UI scraping feasible but untested'),
        ('Threads', 'Video post', 'Firefox + Claude Vision', 'No API; UI scraping feasible but not yet implemented'),
    ],
    [1.2, 1.0, 1.8, 2.5]
)

# ── 9. FUNCTIONAL REQUIREMENTS ────────────────────────────────────────────────
add_heading(doc, '11. Functional Requirements', 1)
add_body(doc, 'Priority: M = Must Have   S = Should Have   C = Could Have', italic=True, color=GREY)
doc.add_paragraph()

M = ('M', False, GREEN)
S = ('S', False, LBLUE)
C = ('C', False, ORANGE)

make_table(doc,
    ['ID', 'Priority', 'Requirement'],
    [
        ('FR-01', M, 'System reads Lark Bitable to find rows where Column F has a URL and columns A/D/E/G are empty.'),
        ('FR-02', M, 'System routes YouTube URLs to yt-dlp for fast, exact metadata extraction (no browser). Routes TikTok and X (Twitter) URLs to Firefox + DOM extraction.'),
        ('FR-03', M, 'System routes Instagram/RedNote URLs to Firefox + Claude Vision for screenshot-based extraction.'),
        ('FR-04', M, 'For Instagram: system takes two screenshots — reel page screenshot (view count, captured immediately before any navigation) and split-view screenshot (caption + date, from clicking the profile grid thumbnail).'),
        ('FR-05', M, 'System sends screenshots to Claude Haiku Vision (claude-haiku-4-5) to extract posted date, caption, view count.'),
        ('FR-06', M, 'System strips ALL hashtags from caption using Unicode-aware regex before writing to Column D.'),
        ('FR-07', M, 'System writes posted date to Column A only.'),
        ('FR-08', M, 'System writes cleaned caption to Column D only. Writes "No caption" if empty.'),
        ('FR-09', M, 'System writes "Content Casual" to Column E only. Hardcoded always.'),
        ('FR-10', M, 'System writes view count to Column G only.'),
        ('FR-11', M, 'System NEVER writes to Columns B, C, F, or H. Any attempt raises an exception immediately.'),
        ('FR-12', M, 'System writes all fields for a row in ONE atomic Lark API call — prevents partial data if system crashes.'),
        ('FR-13', M, 'System writes into existing rows only — never appends new rows.'),
        ('FR-14', M, 'System runs as a background watcher (watcher.py) — checks Lark every 3 minutes automatically.'),
        ('FR-15', M, 'System validates all required .env variables on startup — shows clear error if any are missing.'),
        ('FR-16', M, 'System retries failed Lark writes up to 3 times with exponential back-off (2s, 4s, 8s).'),
        ('FR-17', S, 'System prints a run summary showing success/failure status for each URL processed.'),
        ('FR-18', S, 'If a URL fails, system logs the error in plain English and continues processing remaining URLs.'),
        ('FR-19', S, 'All activity logged to logs/auto_count.log with timestamps and severity levels.'),
        ('FR-20', C, 'System compares current cycle performance against previous cycle.'),
    ],
    [0.8, 0.8, 5.9]
)

# ── 10. NON-FUNCTIONAL REQUIREMENTS ──────────────────────────────────────────
add_heading(doc, '12. Non-Functional Requirements', 1)
make_table(doc,
    ['ID', 'Category', 'Requirement'],
    [
        ('NFR-01', 'Performance', 'Process up to 20 video URLs within 3 minutes.'),
        ('NFR-02', 'Accuracy', 'View counts must match what is shown on the video page within 5% variance (consistent with AC-03). Captions must have zero hashtags remaining after processing (AC-02). Overall field-fill accuracy: 95%+ of submitted URLs result in all 4 columns correctly populated (AC-01).'),
        ('NFR-03', 'Data Integrity', 'System must never write to Columns B, C, F, or H under any circumstance.'),
        ('NFR-04', 'Data Integrity', 'System must never overwrite existing data in columns A, D, E, G if already filled.'),
        ('NFR-05', 'Reliability', 'If one URL fails, system must continue processing all remaining URLs.'),
        ('NFR-06', 'Security', 'API keys stored in .env file only. Never logged or printed. .env excluded from git.'),
        ('NFR-07', 'Testability', '283 automated pytest tests across 6 test files — positive, negative, security, performance, UI/UX, and stress cases. All must pass (no real network calls).'),
        ('NFR-08', 'Compatibility', 'System runs on Python 3.9+ on macOS.'),
        ('NFR-09', 'Observability', 'All activity logged to logs/auto_count.log with timestamps. Errors shown in plain English, not technical jargon.'),
        ('NFR-10', 'Resilience', 'Lark API writes retried up to 3 times on failure. Startup validates all required .env vars before any work begins.'),
    ],
    [0.9, 1.6, 5.0]
)

# ── 11. RISK REGISTER ─────────────────────────────────────────────────────────
add_heading(doc, '13. Risk Register', 1)
H2 = ('High', False, RED)
M2 = ('Medium', False, ORANGE)
L2 = ('Low', False, GREEN)
make_table(doc,
    ['ID', 'Risk', 'Likelihood', 'Impact', 'Mitigation'],
    [
        ('R-01', 'Platform blocks Playwright browser (detects automation)', M2, H2, 'Use persistent Chrome profile with real login session; use slow human-like interactions'),
        ('R-02', 'Claude Vision extracts wrong data from screenshot', L2, H2, 'Validate extracted data format; flag outliers for manual review'),
        ('R-03', 'System writes to wrong Lark column', L2, H2, 'Hard allowlist in lark_writer.py raises exception immediately if wrong column attempted'),
        ('R-04', 'Hashtag regex misses Unicode hashtags (e.g. Chinese, Arabic)', M2, M2, 'Use Unicode-aware regex pattern covering all Unicode word characters'),
        ('R-05', 'RedNote requires login — Playwright session expires', H2, M2, 'Persistent Chrome profile retains login session; user re-logs if expired'),
        ('R-06', 'Lark Sheet API token expires during run', L2, M2, 'Implement token refresh logic; retry on auth failure'),
    ],
    [0.6, 2.3, 1.0, 0.8, 2.8]
)

# ── 12. ASSUMPTIONS & CONSTRAINTS ────────────────────────────────────────────
add_heading(doc, '14. Assumptions & Constraints', 1)

add_heading(doc, '14.1 Assumptions', 2)
for item in [
    'The user pastes video URLs into Column F of the Lark Sheet before running the system.',
    'The user is already logged into all 4 active Phase 1 platforms (Instagram, YouTube, TikTok, RedNote) in the persistent Firefox profile.',
    'The Lark Sheet exists with exactly columns A to H as defined in Section 6.',
    'Videos will always be publicly visible when the browser opens the URL.',
    'The user runs the system on macOS with Python 3.11+ installed.',
    '"Content Casual" is the correct default for all videos in Phase 1.',
]:
    add_bullet(doc, item)

add_heading(doc, '14.2 Constraints', 2)
for item in [
    'System MUST NEVER write to Columns B, C, F, or H — enforced with hard exception in code.',
    'Column F is owned by the user — the system only reads it, never writes to it.',
    'System writes into existing rows only — never creates new rows.',
    'No platform APIs are used — all data is extracted via Playwright + Claude Vision.',
    'Single user only in Phase 1.',
    'System runs via python run.py only — no web UI, no mobile app.',
]:
    add_bullet(doc, item)

# ── 13. IMPLEMENTATION ROADMAP ───────────────────────────────────────────────
add_heading(doc, '15. Implementation Roadmap', 1)
make_table(doc,
    ['Phase', 'What Gets Built', 'Key Output'],
    [
        ('Phase 1\nScaffold', 'Project setup: requirements.txt, .gitignore, .env.template, virtualenv', 'Clean project structure ready to build'),
        ('Phase 2\nLark Writer', 'lark_writer.py with hard column allowlist + pytest tests proving B/C/F/H raise exceptions', 'Green pytest output — safety guaranteed'),
        ('Phase 3\nLark Reader', 'lark_reader.py — find rows where F has URL and A/D/E/G are empty', 'System knows what to process'),
        ('Phase 4\nProcessor', 'processor.py — clean_caption (Unicode hashtag strip), format_date, set_content_type', 'Clean data ready to write'),
        ('Phase 5\nBrowser', 'browser_reader.py — Playwright with persistent Chrome profile', 'Browser opens video pages automatically'),
        ('Phase 6\nVision', 'vision_extract.py — screenshot to Claude Haiku Vision to JSON', 'AI extracts posted_date, caption, view_count'),
        ('Phase 7\nWiring', 'run.py + reporter.py — wire everything together + run summary', 'Single command runs full system'),
        ('Phase 8\nTesting', 'End-to-end test on 3 real YouTube URLs', 'Proven working system'),
    ],
    [1.2, 4.3, 2.0]
)

# ── 14. SUCCESS CRITERIA ─────────────────────────────────────────────────────
add_heading(doc, '16. Acceptance Criteria', 1)
add_body(doc,
    'The system is considered complete and accepted when ALL of the following criteria are met. '
    'Each criterion is specific and measurable — not vague.',
    bold=True
)
doc.add_paragraph()

make_table(doc,
    ['#', 'Acceptance Criterion', 'How to Verify', 'Pass / Fail'],
    [
        ('AC-01',
         'All 4 columns (A, D, E, G) are filled correctly for at least 95% of valid URLs submitted.',
         'Paste 20 real video URLs. Count how many rows have all 4 columns filled correctly in Lark.',
         'PASS if 19+ out of 20 rows are fully correct.'),
        ('AC-02',
         'Zero hashtags remain in Column D (Title) for any processed video.',
         'Check Column D for any word starting with #. Search the entire column.',
         'PASS if 0 hashtags found in any cell.'),
        ('AC-03',
         'View count in Column G matches the number shown on the platform within 5%.',
         'Manually check the view count on platform for 10 processed videos. Compare to Column G.',
         'PASS if all 10 are within 5% of the actual displayed count.'),
        ('AC-04',
         'Posted date in Column A matches the actual video post date.',
         'Check 10 processed videos. Compare Column A date to the date shown on the platform.',
         'PASS if all 10 dates are correct.'),
        ('AC-05',
         'Column F is never modified by the system. Original URL always preserved exactly.',
         'Compare Column F before and after running the system for 20 rows.',
         'PASS if all 20 URLs are byte-for-byte identical after processing.'),
        ('AC-06',
         'Columns B, C, and H are never written to. Any attempt raises an error immediately.',
         'Run pytest test suite — test_lark_writer.py. All 283 automated tests must pass.',
         'PASS if all 283 pytest tests are green.'),
        ('AC-07',
         'Time from pasting URL to data appearing in Lark is under 3 minutes.',
         'Paste a YouTube URL. Start a timer. Check when Lark row is filled.',
         'PASS if data appears within 3 minutes.'),
        ('AC-08',
         'System processes a new URL without the user opening any platform manually.',
         'Paste a URL. Do not open Instagram, YouTube, TikTok, or RedNote. Data should appear automatically.',
         'PASS if Lark fills correctly without user opening any platform.'),
        ('AC-09',
         'If one URL fails (e.g. private video), the system logs the error and continues processing remaining URLs.',
         'Include one private/deleted video URL in a batch of 5. Check that the other 4 are processed.',
         'PASS if 4 rows are filled and 1 error is logged in plain English.'),
        ('AC-10',
         'All error messages are in plain English. No raw Python errors shown to the user.',
         'Trigger common errors (wrong .env, no internet). Read the error message shown.',
         'PASS if all messages are human-readable with a clear action to take.'),
    ],
    [0.5, 2.8, 2.2, 2.0]
)

# ── 15. GLOSSARY ─────────────────────────────────────────────────────────────
add_heading(doc, '17. Glossary', 1)
make_table(doc,
    ['Term', 'Definition'],
    [
        ('Playwright', 'A Python library that controls a real Chrome browser automatically.'),
        ('Claude Vision', 'Claude AI\'s ability to look at a screenshot and extract information from it.'),
        ('Persistent Chrome Profile', 'A saved browser session that remembers your logins across runs.'),
        ('Hard Allowlist', 'A strict list of allowed columns in code — any other column raises an exception immediately.'),
        ('Unicode-aware Regex', 'A pattern that matches hashtags in any language, including Chinese, Arabic, and emoji-based tags.'),
        ('run.py', 'The single Python file the user runs to trigger the entire system.'),
        ('lark_writer.py', 'The module that writes data to Lark Sheet — enforces the column allowlist.'),
        ('lark_reader.py', 'The module that reads the Lark Sheet to find rows needing processing.'),
        ('vision_extract.py', 'The module that sends screenshots to Claude Vision and returns structured data.'),
        ('Content Casual', 'The hardcoded default value always written to Column E.'),
        ('Final Reach', 'Column H — manually calculated by the user. System never touches this.'),
    ],
    [2.0, 5.5]
)

doc.add_paragraph()

# ── 16. SIGN-OFF ──────────────────────────────────────────────────────────────
add_heading(doc, '18. Sign-Off & Approval', 1)
add_body(doc, 'By signing below, the approver confirms that this BRD accurately represents the requirements for the Auto Count Social Media Reach system.')
doc.add_paragraph()
make_table(doc,
    ['Role', 'Name', 'Signature', 'Date'],
    [
        ('Prepared By', 'jingyi-rere', '', ''),
        ('Reviewed By', '', '', ''),
        ('Approved By', '', '', ''),
    ],
    [2.0, 2.0, 2.0, 1.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— End of Document — Version 5.5 —')
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(*GREY)

output_path = '/Users/jjjyyy/Auto-Count Social MediaReach/BRD_Auto_Count_Social_Media_Reach.docx'
doc.save(output_path)
print(f'Saved: {output_path}')

md_path = '/Users/jjjyyy/Auto-Count Social MediaReach/BRD_Auto_Count_Social_Media_Reach.md'
today = datetime.date.today().strftime('%d %B %Y')
md = f"""# Business Requirements Document
## Auto Count Social Media Reach — AI-Assisted Video Metrics Automation System

**Version:** 5.5 | **Date:** {today} | **Status:** Final | **Prepared By:** jingyi-rere

---

## Version History

| Version | Date | Author | Changes |
|---------|------|--------|---------|
| 5.0 | 13 May 2026 | jingyi-rere | Watcher daemon, yt-dlp, Firefox+Vision, 283 tests |
| 5.1 | 13 May 2026 | jingyi-rere | Pain metrics, AI justification, 33 hrs/year value, AC-01–AC-10 |
| 5.2 | 13 May 2026 | jingyi-rere | Phase 1/2 scope clarified, baseline 40 min/week, accuracy thresholds aligned |
| 5.3 | 18 May 2026 | jingyi-rere | Instagram view count from reel screenshot before Escape; watcher 3 min |
| 5.4 | {today} | jingyi-rere | Defensive baseline note removed (stated upfront); Phase 2 moved to Appendix A |

---

## 1. Executive Summary

A video content creator at ricebowlmy spends ~40 minutes every week manually collecting
performance data (posted date, caption, view count) from 4 platforms (Instagram, YouTube,
TikTok, RedNote) and typing it into Lark Bitable. This system eliminates that 40 minutes.

The user pastes video URLs into Lark Column F. A background watcher detects new URLs within
3 minutes and auto-fills: posted date (Column A), caption without hashtags (Column D),
content type (Column E), and view count (Column G).

Claude AI Vision is required — not optional — because Instagram, TikTok, and RedNote have
no public API, block scrapers, and change UI regularly. Only Vision can read any screenshot
regardless of platform, language, or UI version.

**Projected outcome:** 38 minutes saved per week ≈ 33 hours per year, near-zero errors.

---

## 2. Hard Rules (NEVER Violate)

| Rule | Description |
|------|-------------|
| RULE-01 | Only write to Lark columns A, D, E, G. Writing to B, C, F, H raises an exception. |
| RULE-02 | Never overwrite Column F. User pastes URLs there — system never touches it. |
| RULE-03 | Never append new rows. Always write into the existing row. |
| RULE-04 | Strip ALL hashtags from caption (Unicode-aware) before writing to Column D. |
| RULE-05 | Column E must always be exactly "Content Casual" — hardcoded, no exceptions. |

---

## 3. Problem Statement

### 3.1 Time Breakdown (10 videos/week)

| Task (per video) | Time | Specific Problem |
|------------------|------|-----------------|
| Open platform, find the video | ~1 min | Instagram grid paginated; RedNote UI in Chinese |
| Read and copy the caption | ~1 min | Long captions, overlaid text |
| Remove hashtags manually | ~1 min | 15–25 tags in EN/Malay/Chinese; easy to miss hidden tags |
| Note the view count | ~30 sec | Instagram hides view counts on desktop |
| Type data into Lark Bitable | ~30 sec | 4 fields, typos happen (10,000 → 1,000) |
| **TOTAL per video** | **~4 min** | |
| **TOTAL per week (10 videos)** | **~40 min** | 40 min of pure manual data entry every week |

### 3.2 Why Simple Automation Cannot Solve This

- Instagram, TikTok, RedNote have no public API for view counts
- All three platforms actively block headless browsers and scrapers
- UI layouts change without notice — hardcoded selectors break within days
- Content is in multiple languages (EN, Malay, Chinese) — regex extraction fails
- Only AI Vision can read any screenshot regardless of platform or UI version

---

## 4. Why AI (Not Just Automation)

| Approach | Result |
|----------|--------|
| Platform APIs | Instagram and RedNote have no public API for view counts |
| CSS scraping | Blocked by anti-bot measures; selectors change with every UI update |
| yt-dlp (YouTube/TikTok only) | Works perfectly for YouTube and TikTok — used in this system |
| Claude Vision (Instagram/RedNote) | Reads any screenshot regardless of language or UI version |

---

## 5. Quantified Value

| Metric | Value |
|--------|-------|
| Time saved per video | ~4 minutes |
| Videos per week | ~10 |
| Time saved per week | ~38 minutes |
| Weeks per year | 52 |
| **Total hours saved per year** | **~33 hours** |
| Error rate (manual) | ~5% typos/misreads |
| Error rate (automated) | <1% (Vision occasionally misreads) |

---

## 6. Business Objectives

- BO-01: Eliminate manual data entry for 4 platforms
- BO-02: Auto-fill Lark Bitable within 3 minutes of URL paste
- BO-03: Zero risk of overwriting user data (Column F, H protected)
- BO-04: System must work without any platform API key
- BO-05: Maintainable by a non-developer user (run.py one-command start)

---

## 7. User Stories

| ID | As a… | I want… | So that… |
|----|-------|---------|----------|
| US-01 | Content creator | URLs auto-detected when pasted | I don't start a run manually per video |
| US-02 | Content creator | Caption auto-cleaned of hashtags | I can search/filter captions in Lark |
| US-03 | Content creator | View count filled automatically | I save 30 sec per video of platform navigation |
| US-04 | Content creator | Posted date extracted correctly | My Lark timeline stays accurate |
| US-05 | Content creator | System to never touch Column H | My manual calculations are never overwritten |

---

## 8. Lark Sheet Column Mapping

| Column | Field | Written By | Notes |
|--------|-------|-----------|-------|
| A | Date | System | Posted date (YYYY-MM-DD) |
| B | Title | User | NOT touched by system |
| C | Content Type | User | NOT touched by system |
| D | Caption | System | Hashtags stripped |
| E | Content Type | System | Always "Content Casual" |
| F | Link | User | URL pasted by user — NEVER overwritten |
| G | Reach | System | View count as integer |
| H | Final Reach | User | Manual calculation — NEVER touched |

---

## 9. How the System Works

1. User pastes video URL into Column F of Lark Bitable
2. watcher.py polls Lark every 3 minutes for rows where F has URL but A or G is empty
3. For YouTube/TikTok: yt-dlp fetches title, date, view count (no browser needed, <10 sec)
4. For Instagram/RedNote: Firefox loads the page using saved login session
5. Instagram: two-screenshot approach — reel page screenshot (before Escape) for caption+date; profile grid for view count
6. RedNote: single screenshot of video page for all three fields
7. Screenshot sent to Claude Haiku Vision → returns JSON {{posted_date, caption, view_count}}
8. Hashtags stripped with Unicode-aware regex
9. Single atomic Lark API call writes A, D, E, G — with retry (3 attempts, 2s/4s/8s backoff)

---

## 10. Supported Platforms (Current Scope)

| Platform | Method | Notes |
|----------|--------|-------|
| YouTube | yt-dlp | Exact numbers, no browser, <10 sec |
| TikTok | yt-dlp | Exact numbers, no browser, <10 sec |
| Instagram | Firefox + Claude Vision | Two-screenshot: reel page (before Escape) + profile grid |
| RedNote | Firefox + Claude Vision | Single screenshot of video page |

### Appendix A: Future Scope (Not Built)

| Platform | Status | Blocker |
|----------|--------|---------|
| Facebook | Not built | No public API; login wall varies by account type |
| X (Twitter) | Not built | API access requires paid tier |
| Threads | Not built | No API; UI too new/unstable for reliable Vision |

---

## 11. Functional Requirements

| ID | Requirement | Priority |
|----|-------------|----------|
| FR-01 | Detect new URLs in Column F within one watcher cycle (≤3 min) | Must Have |
| FR-02 | Extract posted date, caption, view count for all 4 platforms | Must Have |
| FR-03 | Strip all hashtags from caption before writing | Must Have |
| FR-04 | Write to columns A, D, E, G only | Must Have |
| FR-05 | Never overwrite existing data in A, D, E, G | Must Have |
| FR-06 | Log all activity to logs/auto_count.log | Should Have |
| FR-07 | Retry failed Lark writes up to 3 times | Should Have |
| FR-08 | Display human-readable errors if extraction fails | Should Have |

---

## 12. Non-Functional Requirements

| ID | Requirement | Target |
|----|-------------|--------|
| NFR-01 | Processing time per URL (YouTube/TikTok) | <10 seconds |
| NFR-02 | Processing time per URL (Instagram/RedNote) | <3 minutes |
| NFR-03 | View count accuracy | >99% for yt-dlp; >95% for Vision |
| NFR-04 | System uptime | Runs until manually stopped |
| NFR-05 | Test coverage | 283 automated tests, 0 failures |

---

## 13. Risk Register

| Risk | Likelihood | Impact | Mitigation |
|------|-----------|--------|-----------|
| Instagram UI change | Medium | High | Vision reads screenshots — no CSS selectors to break |
| Vision misreads number | Low | Medium | User spot-checks weekly; value is time-saving not 100% accuracy |
| yt-dlp blocked | Low | High | Fallback: manual entry (unchanged from current process) |
| Lark API rate limit | Low | Medium | Retry with exponential backoff |
| Firefox login session expires | Medium | Medium | User re-logs in manually; session persists via profile |

---

## 16. Acceptance Criteria

| ID | Criterion | Pass Condition |
|----|-----------|---------------|
| AC-01 | URL detection | New URL processed within 3 min of paste |
| AC-02 | Date accuracy | Matches platform display date |
| AC-03 | View count accuracy | Within 1% for yt-dlp; within 5% for Vision |
| AC-04 | Caption accuracy | Exact caption minus hashtags |
| AC-05 | Column protection | Columns B, C, F, H never written |
| AC-06 | Hashtag removal | Zero hashtags remain in Column D |
| AC-07 | Content type | Column E always "Content Casual" |
| AC-08 | No duplicate writes | Row with existing A+G skipped |
| AC-09 | Error logging | All failures written to log file |
| AC-10 | PIC filter | Only processes rows where PIC = "TAN JING YI" |

---

*End of Document — Version 5.4*
"""
with open(md_path, 'w', encoding='utf-8') as f:
    f.write(md)
print(f'Saved: {md_path}')
